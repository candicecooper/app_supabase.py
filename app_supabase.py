import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime, date, time, timedelta
import uuid
import random
from io import BytesIO
import base64
import bcrypt

# SUPABASE CONNECTION
try:
    from supabase import create_client, Client
    SUPABASE_AVAILABLE = True
except ImportError:
    SUPABASE_AVAILABLE = False
    st.warning("Supabase not installed. Run: pip install supabase")

# Initialize Supabase client
@st.cache_resource
def init_supabase() -> Client:
    """Initialize Supabase client with credentials from secrets"""
    if not SUPABASE_AVAILABLE:
        return None
    
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
        return create_client(url, key)
    except Exception as e:
        st.error(f"‚ùå Supabase connection failed: {e}")
        st.info("üí° Add Supabase credentials to .streamlit/secrets.toml")
        return None

# Global Supabase client
supabase: Client = init_supabase()

st.set_page_config(page_title="CLC Behaviour Support", page_icon="üìä", layout="wide", initial_sidebar_state="collapsed")

# MINIMALIST PROFESSIONAL STYLING
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    * { font-family: 'Inter', sans-serif; }
    .stApp { background: #f8fafc; }
    .stButton>button {
        background: #334155 !important; color: white !important;
        border: none !important; border-radius: 6px !important;
        padding: 0.5rem 1.2rem !important; font-weight: 600 !important;
        box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        transition: all 0.2s !important;
    }
    .stButton>button:hover { background: #1e293b !important; transform: translateY(-1px) !important; }
    button[kind="primary"] { background: #0ea5e9 !important; color: white !important; }
    button[kind="primary"]:hover { background: #0284c7 !important; }
    [data-testid="stVerticalBlock"] > div[style*="border"] {
        background: white !important; border-radius: 8px !important;
        padding: 1.5rem !important; box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
        border: 1px solid #e2e8f0 !important;
    }
    [data-testid="stMetricValue"] {
        font-size: 1.875rem !important; font-weight: 700 !important; color: #0f172a !important;
    }
    [data-testid="stMetricLabel"] {
        color: #64748b !important; font-weight: 600 !important; font-size: 0.875rem !important;
        text-transform: uppercase !important; letter-spacing: 0.05em !important;
    }
    .stTextInput>div>div>input, .stSelectbox>div>div>select, 
    .stTextArea>div>div>textarea, .stNumberInput>div>div>input,
    .stDateInput>div>div>input, .stTimeInput>div>div>input {
        border: 1px solid #cbd5e1 !important; background: white !important;
        color: #0f172a !important; font-weight: 500 !important; border-radius: 6px !important;
    }
    h1 { color: #0f172a !important; font-weight: 700 !important; }
    h2 { color: #0f172a !important; font-weight: 700 !important; }
    h3 { color: #0f172a !important; font-weight: 600 !important; }
    label { color: #334155 !important; font-weight: 600 !important; font-size: 0.875rem !important; }
    .stSuccess { background: #ecfdf5 !important; color: #065f46 !important; 
                 border-left: 4px solid #10b981 !important; }
    .stInfo { background: #f0f9ff !important; color: #075985 !important; 
              border-left: 4px solid #0ea5e9 !important; }
    .stWarning { background: #fffbeb !important; color: #92400e !important; 
                 border-left: 4px solid #f59e0b !important; }
</style>
""", unsafe_allow_html=True)

# Production mode - sandbox banner removed

# MOCK DATA
MOCK_STAFF = [
    {"id": "s1", "name": "Emily Jones", "role": "JP", "email": "emily.jones@example.com", "password": "demo123"},
    {"id": "s2", "name": "Daniel Lee", "role": "PY", "email": "daniel.lee@example.com", "password": "demo123"},
    {"id": "s3", "name": "Sarah Chen", "role": "SY", "email": "sarah.chen@example.com", "password": "demo123"},
    {"id": "s4", "name": "Admin User", "role": "ADM", "email": "admin@example.com", "password": "admin123"},
]

MOCK_STUDENTS = [
    {"id": "stu_jp1", "name": "Emma T.", "grade": "R", "dob": "2018-05-30", "program": "JP", 
     "edid": "ED123456", "placement_start": "2024-02-01", "placement_end": None},
    {"id": "stu_jp2", "name": "Oliver S.", "grade": "Y1", "dob": "2017-09-12", "program": "JP",
     "edid": "ED234567", "placement_start": "2024-03-15", "placement_end": None},
    {"id": "stu_jp3", "name": "Sophie M.", "grade": "Y2", "dob": "2016-03-20", "program": "JP",
     "edid": "ED345678", "placement_start": "2024-01-29", "placement_end": None},
    {"id": "stu_py1", "name": "Liam C.", "grade": "Y3", "dob": "2015-06-15", "program": "PY",
     "edid": "ED456789", "placement_start": "2024-02-12", "placement_end": None},
    {"id": "stu_py2", "name": "Ava R.", "grade": "Y4", "dob": "2014-11-08", "program": "PY",
     "edid": "ED567890", "placement_start": "2024-01-08", "placement_end": None},
    {"id": "stu_py3", "name": "Noah B.", "grade": "Y6", "dob": "2012-02-28", "program": "PY",
     "edid": "ED678901", "placement_start": "2024-04-03", "placement_end": None},
    {"id": "stu_sy1", "name": "Isabella G.", "grade": "Y7", "dob": "2011-04-17", "program": "SY",
     "edid": "ED789012", "placement_start": "2024-01-29", "placement_end": None},
    {"id": "stu_sy2", "name": "Ethan D.", "grade": "Y9", "dob": "2009-12-03", "program": "SY",
     "edid": "ED890123", "placement_start": "2024-02-26", "placement_end": None},
    {"id": "stu_sy3", "name": "Mia A.", "grade": "Y11", "dob": "2007-08-20", "program": "SY",
     "edid": "ED901234", "placement_start": "2024-03-11", "placement_end": None},
]

PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}
BEHAVIOUR_TYPES = ["Verbal Refusal", "Elopement", "Property Destruction", "Aggression (Peer)", 
                   "Aggression (Adult)", "Self-Harm", "Verbal Aggression", "Other"]
ANTECEDENTS = [
    "--- PEER ---",
    "Peer - negative peer feedback",
    "Peer - peer conflict/interaction", 
    "Peer - participating in competition",
    "Peer - losing against a peer/s",
    "Peer - another student escalating",
    "--- TRANSITION ---",
    "Transition - to a non-preferred activity",
    "Transition - from one activity to another",
    "Transition - from one environment to another",
    "Transition - another teacher coming into the program",
    "Transition - from play/yard to classroom",
    "Transition - from home to the program",
    "Transition - off a device",
    "--- INSTRUCTIONS ---",
    "Instructions - following instructions given by an adult",
    "Instructions - following task demands",
    "--- ENGAGEMENT ---",
    "Engagement - engaging in a non-preferred learning area (Maths/Literacy)",
    "Engagement - having to wait",
    "Engagement - change in routine",
    "Engagement - learning task too difficult or perceived as difficult",
    "Engagement - having to work with a peer",
    "Engagement - working independently",
    "Engagement - interrupted activity",
    "Engagement - not being able to finish an activity",
    "--- SENSORY ---",
    "Sensory - crowded area",
    "Sensory - noise",
    "Sensory - environment too bright",
    "Sensory - no medication",
    "Sensory - hungry/thirsty",
    "Sensory - unable to take a movement break",
    "--- OTHER ---",
    "Other - perceived injustice",
    "Other - TRT (Temporary Relief Teacher)",
    "Other - NIT teacher (New Initiatives Teacher)",
    "Other - incident before coming to program",
    "Other - verbalising not wanting to be at the program",
    "Other - staff attention shifted",
    "Other - unstructured time",
    "Other - access denied"
]
INTERVENTIONS = ["CPI Supportive stance", "Offered break", "Reduced demand", "Provided choices", 
                "Removed audience", "Visual supports", "Co-regulation", "Prompted coping skill", "Redirection"]
LOCATIONS = ["JP Classroom", "PY Classroom", "SY Classroom", "Playground", "Library", "Office", "Student Gate", "Toilets"]
VALID_PAGES = ["login", "landing", "program_students", "incident_log", "critical_incident", "student_analysis", "admin_portal"]

# AI HYPOTHESIS SYSTEM
HYPOTHESIS_FUNCTIONS = ["To get", "To avoid"]
HYPOTHESIS_ITEMS = ["Tangible", "Activity", "Sensory", "Attention"]

# AUSTRALIAN CURRICULUM GENERAL CAPABILITIES
# These capabilities are developed across all learning areas
AC_CAPABILITIES = {
    "PSC": {
        "name": "Personal and Social Capability",
        "elements": {
            "Self-awareness": ["Recognise emotions", "Recognise personal qualities and achievements", "Understand themselves as learners", "Develop reflective practice"],
            "Self-management": ["Express emotions appropriately", "Develop self-discipline and set goals", "Work independently and show initiative", "Become confident, resilient and adaptable"],
            "Social awareness": ["Appreciate diverse perspectives", "Contribute to civil society", "Understand relationships"],
            "Social management": ["Communicate effectively", "Work collaboratively", "Make decisions", "Negotiate and resolve conflict", "Develop leadership skills"]
        },
        "color": "#4A90A4"
    },
    "CCT": {
        "name": "Critical and Creative Thinking",
        "elements": {
            "Inquiring": ["Identify, explore and organise information and ideas", "Pose questions"],
            "Generating": ["Generate ideas, possibilities and actions", "Consider alternatives"],
            "Analysing": ["Analyse, synthesise and evaluate reasoning and procedures"],
            "Reflecting": ["Reflect on thinking and processes", "Apply logic and reasoning"]
        },
        "color": "#6BB9A0"
    },
    "EU": {
        "name": "Ethical Understanding",
        "elements": {
            "Understanding": ["Recognise ethical concepts", "Explore ethical issues"],
            "Reasoning": ["Reason and make ethical decisions"],
            "Acting": ["Consider consequences", "Reflect on ethical action"]
        },
        "color": "#E8B960"
    },
    "ICU": {
        "name": "Intercultural Understanding",
        "elements": {
            "Recognising": ["Recognise culture and develop respect"],
            "Interacting": ["Interact and empathise with others"],
            "Reflecting": ["Reflect on intercultural experiences and take responsibility"]
        },
        "color": "#D4A574"
    }
}

# Behaviour to AC Capability Mapping
BEHAVIOUR_AC_MAPPING = {
    "Verbal Refusal": {
        "primary": "PSC",
        "elements": ["Self-management", "Social management"],
        "skills_to_develop": [
            "Express emotions appropriately",
            "Communicate effectively",
            "Negotiate and resolve conflict"
        ],
        "ac_descriptors": [
            "Persist in the face of difficulty",
            "Express feelings and opinions appropriately",
            "Work toward shared goals"
        ]
    },
    "Elopement": {
        "primary": "PSC",
        "elements": ["Self-awareness", "Self-management"],
        "skills_to_develop": [
            "Recognise emotions",
            "Develop self-discipline and set goals",
            "Become confident, resilient and adaptable"
        ],
        "ac_descriptors": [
            "Identify and express a range of emotions",
            "Identify personal strengths and challenges",
            "Persist when faced with challenges"
        ]
    },
    "Property Destruction": {
        "primary": "PSC",
        "elements": ["Self-management", "Social awareness"],
        "skills_to_develop": [
            "Express emotions appropriately",
            "Understand relationships",
            "Appreciate diverse perspectives"
        ],
        "ac_descriptors": [
            "Control impulses and reactions",
            "Consider points of view of others",
            "Identify the effects of actions on others"
        ]
    },
    "Aggression (Peer)": {
        "primary": "PSC",
        "elements": ["Social management", "Social awareness"],
        "skills_to_develop": [
            "Negotiate and resolve conflict",
            "Communicate effectively",
            "Understand relationships"
        ],
        "ac_descriptors": [
            "Use problem-solving skills to resolve conflict",
            "Consider points of view of others",
            "Develop strategies to manage conflict"
        ]
    },
    "Aggression (Adult)": {
        "primary": "PSC",
        "elements": ["Self-management", "Social management"],
        "skills_to_develop": [
            "Express emotions appropriately",
            "Communicate effectively",
            "Develop self-discipline and set goals"
        ],
        "ac_descriptors": [
            "Express strong emotions appropriately",
            "Develop positive relationships with adults",
            "Respond appropriately to guidance"
        ]
    },
    "Self-Harm": {
        "primary": "PSC",
        "elements": ["Self-awareness", "Self-management"],
        "skills_to_develop": [
            "Recognise emotions",
            "Become confident, resilient and adaptable",
            "Develop reflective practice"
        ],
        "ac_descriptors": [
            "Identify and express emotions in safe ways",
            "Develop coping strategies",
            "Seek help when needed"
        ]
    },
    "Verbal Aggression": {
        "primary": "PSC",
        "elements": ["Self-management", "Social management"],
        "skills_to_develop": [
            "Express emotions appropriately",
            "Communicate effectively",
            "Negotiate and resolve conflict"
        ],
        "ac_descriptors": [
            "Use respectful language",
            "Express disagreement appropriately",
            "Consider impact of words on others"
        ]
    },
    "Other": {
        "primary": "PSC",
        "elements": ["Self-awareness", "Self-management"],
        "skills_to_develop": [
            "Recognise emotions",
            "Express emotions appropriately",
            "Develop self-discipline and set goals"
        ],
        "ac_descriptors": [
            "Identify triggers and patterns",
            "Develop self-regulation strategies",
            "Set and work toward personal goals"
        ]
    }
}

# Antecedent to AC Skills Connection
ANTECEDENT_AC_SKILLS = {
    "Peer": {
        "capability": "PSC",
        "focus_elements": ["Social awareness", "Social management"],
        "teaching_priority": ["Perspective-taking", "Conflict resolution", "Emotional regulation in social contexts"]
    },
    "Transition": {
        "capability": "PSC", 
        "focus_elements": ["Self-management", "Self-awareness"],
        "teaching_priority": ["Flexibility", "Coping with change", "Self-regulation"]
    },
    "Instructions": {
        "capability": "PSC",
        "focus_elements": ["Self-management", "Social management"],
        "teaching_priority": ["Following routines", "Responding to guidance", "Task completion"]
    },
    "Engagement": {
        "capability": "CCT",
        "focus_elements": ["Reflecting", "Generating"],
        "teaching_priority": ["Persistence", "Problem-solving", "Growth mindset"]
    },
    "Sensory": {
        "capability": "PSC",
        "focus_elements": ["Self-awareness", "Self-management"],
        "teaching_priority": ["Body awareness", "Self-advocacy", "Regulation strategies"]
    },
    "Other": {
        "capability": "PSC",
        "focus_elements": ["Self-awareness", "Self-management"],
        "teaching_priority": ["Emotional literacy", "Help-seeking", "Coping strategies"]
    }
}

# CONSISTENT COLOUR SCHEME FOR GRAPHS
CHART_COLORS = {
    "primary": "#008080",      # Teal
    "secondary": "#228B22",    # Forest Green  
    "accent": "#4682B4",       # Steel Blue
    "warning": "#DC3545",      # Red
    "success": "#28A745",      # Green
    "neutral": "#6C757D",      # Gray
    "light": "#E8F4F8",        # Light teal
    "regular_incident": "#4A90A4",   # Teal-blue for regular incidents
    "critical_incident": "#DC3545",  # Red for critical incidents
    "gradient": ["#008080", "#20B2AA", "#48D1CC", "#7FFFD4"]  # Teal gradient
}

def format_time_12hr(time_str):
    """Convert 24hr time string to 12hr format"""
    try:
        if isinstance(time_str, str):
            dt = datetime.strptime(time_str, "%H:%M:%S")
        else:
            dt = datetime.combine(date.today(), time_str)
        return dt.strftime("%I:%M %p")
    except:
        return time_str

def format_date_dmy(date_str):
    """Convert date to DD/MM/YYYY format"""
    try:
        if isinstance(date_str, str):
            dt = datetime.strptime(date_str, "%Y-%m-%d")
        else:
            dt = date_str
        return dt.strftime("%d/%m/%Y")
    except:
        return date_str

def generate_hypothesis(antecedent, behaviour, consequence):
    """Auto-generate hypothesis based on ABC data"""
    hypotheses = []
    antecedent_lower = antecedent.lower()
    behaviour_lower = behaviour.lower()
    
    # Updated to match new antecedent categories
    if any(word in antecedent_lower for word in [
        "instruction", "demand", "task", "transition", "work", "non-preferred",
        "literacy", "maths", "wait", "routine", "independently", "difficult"
    ]):
        hypotheses.append("To avoid or escape the demand/task")
        
    if any(word in antecedent_lower for word in [
        "attention", "shifted", "ignored", "alone", "feedback", "conflict",
        "interaction", "escalating", "injustice"
    ]):
        hypotheses.append("To gain staff/peer attention")
        
    if any(word in antecedent_lower for word in [
        "sensory", "loud", "noise", "bright", "touch", "crowded", "medication",
        "hungry", "thirsty", "movement"
    ]):
        hypotheses.append("To escape sensory discomfort or seek sensory input")
        
    if any(word in antecedent_lower for word in [
        "denied", "can't have", "no", "wait", "device", "finish", "preferred"
    ]):
        hypotheses.append("To gain access to preferred item/activity")
        
    if any(word in behaviour_lower for word in ["refusal", "defiance", "left", "ran", "elopement"]):
        hypotheses.append("To assert control or autonomy")
    
    if not hypotheses:
        hypotheses.append("Function requires further analysis")
    
    return " / ".join(hypotheses[:2])

def generate_admin_summary(incident_data, student, staff_name):
    """Generate brief summary for external incident log - FOR ADMIN USE ONLY"""
    abch_primary = incident_data.get("ABCH_primary", {})
    intended = incident_data.get("intended_outcomes", [])
    
    date_time = incident_data.get("created_at", datetime.now().isoformat())
    dt = datetime.fromisoformat(date_time)
    
    location = abch_primary.get("location", "Unknown location")
    time_str = abch_primary.get("time", "Unknown time")
    behaviour = abch_primary.get("behaviour", "Behaviour not specified")
    context = abch_primary.get("context", "")
    consequence = abch_primary.get("consequence", "")
    hypothesis = abch_primary.get("hypothesis", "")
    
    summary = f"""CRITICAL INCIDENT SUMMARY - FOR ADMIN USE ONLY
    
Student: {student['name']} (Grade {student['grade']})
Date/Time: {dt.strftime('%d/%m/%Y')} at {time_str}
Location: {location}

INCIDENT DESCRIPTION:
{behaviour[:200]}{'...' if len(behaviour) > 200 else ''}

CONTEXT/ANTECEDENT:
{context[:200]}{'...' if len(context) > 200 else ''}

IMMEDIATE STAFF RESPONSE:
{consequence[:200]}{'...' if len(consequence) > 200 else ''}

BEHAVIOURAL FUNCTION:
{hypothesis}

OUTCOMES IMPLEMENTED:
{', '.join(intended[:5]) if intended else 'See full form for details'}

REPORTED BY: {staff_name}
FORM COMPLETED: {dt.strftime('%d/%m/%Y %H:%M')}

** This summary is for external departmental incident log entry purposes only **
** Full critical incident form has been saved and distributed to relevant parties **
"""
    
    return summary


def generate_hypothesis_ai(antecedent, behaviour, consequence=""):
    """AI generates structured hypothesis from ABC data - ENHANCED VERSION"""
    ant_lower = (antecedent or "").lower()
    beh_lower = (behaviour or "").lower()
    cons_lower = (consequence or "").lower()
    
    # Expanded keyword detection
    avoid_keywords = [
        "instruction", "demand", "task", "transition", "work", "difficult", "challenging",
        "non-preferred", "literacy", "maths", "wait", "routine", "peer", "independently",
        "interrupted", "finish", "teacher coming"
    ]
    
    get_keywords = [
        "attention", "item", "toy", "want", "access", "denied", "device", "finish",
        "preferred", "competition", "feedback"
    ]
    
    sensory_keywords = [
        "sensory", "loud", "noise", "touch", "bright", "crowded", "medication",
        "hungry", "thirsty", "movement break"
    ]
    
    attention_keywords = [
        "attention", "staff", "peer", "ignored", "escalating", "feedback",
        "conflict", "interaction", "injustice"
    ]
    
    # Determine FUNCTION (To get vs To avoid)
    function = "To avoid"
    if any(word in ant_lower for word in get_keywords):
        function = "To get"
    elif any(word in cons_lower for word in ["given", "received", "got", "obtained"]):
        function = "To get"
    elif "denied" in ant_lower or "can't" in ant_lower or "wait" in ant_lower:
        function = "To get"
    elif "off a device" in ant_lower or "finish" in ant_lower:
        function = "To get"
    
    # Determine ITEM (What they want to get or avoid)
    item = "Activity"
    
    if any(word in ant_lower + beh_lower for word in attention_keywords):
        item = "Attention"
    elif any(word in ant_lower + beh_lower for word in sensory_keywords):
        item = "Sensory"
    elif any(word in ant_lower + beh_lower for word in ["toy", "item", "object", "food", "device", "tangible"]):
        item = "Tangible"
    elif any(word in ant_lower for word in ["instruction", "demand", "task", "work", "learning", "maths", "literacy"]):
        item = "Activity"
    elif "transition" in ant_lower or "environment" in ant_lower or "classroom" in ant_lower:
        item = "Activity"
    
    return {"function": function, "item": item}

def format_hypothesis(hyp):
    """Format hypothesis dict to string"""
    if isinstance(hyp, dict):
        return f"{hyp.get('function', 'Unknown')} {hyp.get('item', 'Unknown')}"
    elif isinstance(hyp, str):
        return hyp
    else:
        return "Unknown"

def get_ac_capability_for_behaviour(behaviour_type):
    """Get Australian Curriculum capability information for a behaviour type"""
    if behaviour_type in BEHAVIOUR_AC_MAPPING:
        mapping = BEHAVIOUR_AC_MAPPING[behaviour_type]
        capability_code = mapping["primary"]
        capability_info = AC_CAPABILITIES.get(capability_code, {})
        return {
            "code": capability_code,
            "name": capability_info.get("name", ""),
            "elements": mapping["elements"],
            "skills_to_develop": mapping["skills_to_develop"],
            "ac_descriptors": mapping["ac_descriptors"],
            "color": capability_info.get("color", "#4A90A4")
        }
    return None

def get_ac_skills_for_antecedent(antecedent):
    """Get AC skill focus based on antecedent category"""
    # Determine antecedent category from the full antecedent string
    ant_lower = antecedent.lower() if antecedent else ""
    
    category = "Other"
    if "peer" in ant_lower:
        category = "Peer"
    elif "transition" in ant_lower:
        category = "Transition"
    elif "instruction" in ant_lower:
        category = "Instructions"
    elif "engagement" in ant_lower or "demand" in ant_lower or "task" in ant_lower:
        category = "Engagement"
    elif "sensory" in ant_lower:
        category = "Sensory"
    
    if category in ANTECEDENT_AC_SKILLS:
        skills = ANTECEDENT_AC_SKILLS[category]
        capability_info = AC_CAPABILITIES.get(skills["capability"], {})
        return {
            "category": category,
            "capability_code": skills["capability"],
            "capability_name": capability_info.get("name", ""),
            "focus_elements": skills["focus_elements"],
            "teaching_priority": skills["teaching_priority"],
            "color": capability_info.get("color", "#4A90A4")
        }
    return None

def generate_ac_learning_goals(behaviour_type, antecedent, grade):
    """Generate AC-aligned learning goals based on behaviour and grade level"""
    beh_mapping = BEHAVIOUR_AC_MAPPING.get(behaviour_type, BEHAVIOUR_AC_MAPPING["Other"])
    
    # Adjust complexity based on grade
    try:
        grade_num = int(grade.replace("Y", "").replace("R", "0"))
    except:
        grade_num = 3  # Default to middle primary
    
    # Generate grade-appropriate goals
    if grade_num <= 2:  # Foundation to Year 2
        complexity = "foundational"
        verbs = ["identify", "recognise", "begin to use", "with support"]
    elif grade_num <= 4:  # Years 3-4
        complexity = "developing"
        verbs = ["describe", "demonstrate", "use", "with guidance"]
    elif grade_num <= 6:  # Years 5-6
        complexity = "consolidating"
        verbs = ["explain", "apply", "independently use", "evaluate"]
    else:  # Years 7+
        complexity = "extending"
        verbs = ["analyse", "critically reflect", "independently demonstrate", "adapt"]
    
    goals = []
    for i, skill in enumerate(beh_mapping["skills_to_develop"][:3]):
        verb = verbs[i % len(verbs)]
        goals.append({
            "skill": skill,
            "goal": f"Student will {verb} {skill.lower()}",
            "ac_descriptor": beh_mapping["ac_descriptors"][i] if i < len(beh_mapping["ac_descriptors"]) else "",
            "complexity": complexity
        })
    
    return goals

def get_intervention_ac_alignment(interventions):
    """Map interventions to AC capability development"""
    intervention_ac_map = {
        "CPI Supportive stance": {
            "capability": "PSC",
            "element": "Relationship/Self-management",
            "supports": "Co-regulation and safety"
        },
        "Offered break": {
            "capability": "PSC",
            "element": "Self-management",
            "supports": "Self-regulation and body awareness"
        },
        "Reduced demand": {
            "capability": "PSC",
            "element": "Self-management",
            "supports": "Building stamina gradually"
        },
        "Provided choices": {
            "capability": "PSC",
            "element": "Self-management/Social management",
            "supports": "Agency and decision-making"
        },
        "Removed audience": {
            "capability": "PSC",
            "element": "Social awareness",
            "supports": "Dignity and self-awareness"
        },
        "Visual supports": {
            "capability": "CCT",
            "element": "Analysing/Reflecting",
            "supports": "Processing and understanding expectations"
        },
        "Co-regulation": {
            "capability": "PSC",
            "element": "Self-management/Social awareness",
            "supports": "Emotional regulation through relationship"
        },
        "Prompted coping skill": {
            "capability": "PSC",
            "element": "Self-management",
            "supports": "Building independent regulation"
        },
        "Redirection": {
            "capability": "CCT",
            "element": "Generating/Analysing",
            "supports": "Flexible thinking and alternatives"
        }
    }
    
    alignments = []
    if isinstance(interventions, list):
        for intervention in interventions:
            if intervention in intervention_ac_map:
                alignments.append({
                    "intervention": intervention,
                    **intervention_ac_map[intervention]
                })
    return alignments

def show_severity_guide():
    """Enhanced Behaviour Severity Continuum matching uploaded image"""
    import streamlit.components.v1 as components
    
    html_content = """<div style='background: white; padding: 1.5rem; border-radius: 8px; margin: 1rem 0; 
            box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1); border: 1px solid #e2e8f0;
            font-family: "Source Sans Pro", sans-serif;'>
    
    <div style='text-align: center; margin-bottom: 1.5rem;'>
        <h2 style='color: #1a1a1a; font-weight: 700; font-size: 1.8rem; margin: 0;'>
            Behaviour Severity Continuum
        </h2>
    </div>
    
    <div style='display: grid; grid-template-columns: repeat(5, 1fr); gap: 0;'>
        
        <div style='background: #81b29a; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 1</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Low Level /<br>Engaged
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Following instructions</li>
                    <li>On task</li>
                    <li>Minor defiance</li>
                    <li>Avoiding work</li>
                    <li>Answering back</li>
                    <li>Mumbling, huffing</li>
                    <li>Passive peer conflict</li>
                    <li>Attention seeking</li>
                </ul>
            </div>
            <div style='background: #6b9b7f; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Redirect / provide options</li>
                <li>Give space</li>
                <li>Offer choice</li>
                <li>Acknowledge concern</li>
                <li>Maintain routine</li>
                <li>Active listening</li>
            </ul>
        </div>
        
        <div style='background: #f4d35e; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: #2c2c2c; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 2</div>
                <div style='color: #2c2c2c; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Escalating /<br>Dysregulated
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.3); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: #2c2c2c; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: #2c2c2c; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Raised voice</li>
                    <li>Arguing, blaming</li>
                    <li>Crying, frustration</li>
                    <li>Pacing or mild exit attempts</li>
                    <li>Throwing soft items (not dangerous)</li>
                </ul>
            </div>
            <div style='background: #d9b84d; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: #2c2c2c; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: #2c2c2c; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Reduce demands</li>
                <li>Offer space / movement break</li>
                <li>Provide limited choices</li>
                <li>Avoid power struggles</li>
            </ul>
        </div>
        
        <div style='background: #ee8434; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 3</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    High Escalation /<br>Significant Risk
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Yelling, swearing</li>
                    <li>Slammed doors, hitting walls</li>
                    <li>Throwing items with possible risk</li>
                    <li>Attempting to run off</li>
                    <li>Damaging property</li>
                </ul>
            </div>
            <div style='background: #d47230; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Increase distance</li>
                <li>Notify leadership/support</li>
                <li>Remove audience</li>
                <li>Complete Critical Incident Form</li>
            </ul>
        </div>
        
        <div style='background: #c9555e; padding: 1.2rem 0.8rem; border-right: 2px solid white;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Level 4</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Dangerous<br>Behaviour
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.2); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Attempts to hit, kick, grab</li>
                    <li>Throwing dangerous objects</li>
                    <li>Threats of violence</li>
                    <li>Absconding into unsafe situations</li>
                    <li>Beginning self-harm behaviour</li>
                </ul>
            </div>
            <div style='background: #b04850; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Evacuate nearby students</li>
                <li>Leadership/response team activated</li>
                <li>Maintain safety distance</li>
            </ul>
        </div>
        
        <div style='background: #7d2e2e; padding: 1.2rem 0.8rem;'>
            <div style='text-align: center; margin-bottom: 1rem;'>
                <div style='color: white; font-weight: 700; font-size: 1.1rem; margin-bottom: 0.3rem;'>Crisis</div>
                <div style='color: white; font-weight: 600; font-size: 0.9rem; line-height: 1.3;'>
                    Crisis<br>Situation
                </div>
            </div>
            <div style='background: rgba(255,255,255,0.15); padding: 0.8rem; border-radius: 4px; margin-bottom: 0.8rem;'>
                <div style='color: white; font-weight: 600; font-size: 0.75rem; margin-bottom: 0.4rem;'>Examples:</div>
                <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                    <li>Physical violence causing or likely to cause injury</li>
                    <li>Severe self-harm</li>
                    <li>Use of weapons or dangerous items</li>
                    <li>Full loss of control behaviour</li>
                </ul>
            </div>
            <div style='background: #5c2323; padding: 0.6rem; border-radius: 4px; margin-bottom: 0.5rem;'>
                <div style='color: white; font-weight: 700; font-size: 0.75rem; text-align: center;'>Teacher Priority</div>
            </div>
            <ul style='color: white; font-size: 0.7rem; margin: 0; padding-left: 1.2rem; line-height: 1.5;'>
                <li>Immediate emergency response</li>
                <li>Trained staff to manage situation</li>
                <li>Preserve evidence</li>
                <li>Complete Critical Incident Form</li>
            </ul>
        </div>
        
    </div>

    <div style='margin-top: 1.5rem; padding: 1rem; background: #fff3cd; border-radius: 6px; border-left: 4px solid #f59e0b;'>
        <div style='color: #92400e; font-weight: 700; font-size: 0.95rem; margin-bottom: 0.5rem;'>
            WARNING: Critical Incident Documentation Required
        </div>
        <div style='color: #92400e; font-size: 0.85rem; line-height: 1.5;'>
            <strong>Level 3 or above</strong> requires a Critical Incident ABCH Form to be completed immediately after the incident is resolved.
        </div>
    </div>
</div>
"""
    components.html(html_content, height=700, scrolling=True)
    
def send_critical_incident_email(incident_data, student, staff_email, leader_email, admin_email):
    """Send email notification to all parties"""
    st.info(f"""üìß **Email Notification Sent**
    
**To:** {leader_email}, {admin_email}, {staff_email}  
**Subject:** CRITICAL INCIDENT - {student['name']}

**Student:** {student['name']} | **Programme:** {student['program']} | **Grade:** {student['grade']}  

Critical Incident Form completed and saved.
Admin summary included for departmental log.

*(In production, this sends via SMTP)*
    """)



def generate_behaviour_analysis_plan_docx(student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level):
    """Generate comprehensive BAP with matplotlib graphs (no Chrome/Kaleido needed)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import matplotlib.pyplot as plt
        import matplotlib
        matplotlib.use('Agg')
        
        # ARIAL FONT SETUP
        from docx.oxml.ns import qn
        
        def set_arial(run):
            """Set Arial font for a run"""
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # GREEN COLOR for headings
        GREEN_RGB = RGBColor(34, 139, 34)
        
        
        PROGRAM_NAMES = {"JP": "Junior Primary", "PY": "Primary Years", "SY": "Senior Years"}
        
        doc = Document()
        
        # Set default font to Arial
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
       # ====================================================================
        # BEHAVIOUR ANALYSIS REPORT GENERATION - COMPLETE UPDATED VERSION
        # This replaces lines 337-816 in app_supabase.py
        # ====================================================================
        
        # TITLE PAGE
        heading = doc.add_heading('Behaviour Analysis Plan', 0)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Evidence-Based Analysis & Recommendations')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 116, 139)
        
        doc.add_paragraph()
        
        # Add analysis image
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            import matplotlib.patches as mpatches
            from matplotlib.patches import FancyBboxPatch, Circle
            import numpy as np
            
            fig, ax = plt.subplots(figsize=(6, 3), dpi=150)
            ax.set_xlim(0, 10)
            ax.set_ylim(0, 5)
            ax.axis('off')
            
            # Background
            ax.add_patch(FancyBboxPatch((0, 0), 10, 5, boxstyle="round,pad=0.1", 
                                        facecolor='#f8fafc', edgecolor='#e2e8f0', linewidth=2))
            
            # Bar chart
            bars_x = [1.5, 2.5, 3.5, 4.5, 5.5]
            bars_y = [2.5, 3.2, 2.8, 3.5, 3.0]
            for x, y in zip(bars_x, bars_y):
                ax.add_patch(plt.Rectangle((x-0.3, 0.5), 0.6, y-0.5, 
                                          facecolor='#3b82f6', alpha=0.7))
            
            # Trend line
            line_x = np.linspace(6.5, 9.5, 50)
            line_y = 1.5 + (line_x - 6.5) * 0.2
            ax.plot(line_x, line_y, color='#22c55e', linewidth=3, alpha=0.8)
            ax.scatter([6.5, 7.5, 8.5, 9.5], [1.5, 1.7, 2.1, 2.3], 
                      s=60, color='#22c55e', zorder=5, alpha=0.8)
            
            # Icons
            circle = Circle((1, 1.5), 0.4, facecolor='none', edgecolor='#0ea5e9', linewidth=3)
            ax.add_patch(circle)
            ax.plot([1.3, 1.6], [1.2, 0.9], color='#0ea5e9', linewidth=3)
            
            plt.tight_layout()
            
            img_stream = BytesIO()
            plt.savefig(img_stream, format='png', dpi=150, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            img_stream.seek(0)
            plt.close()
            
            doc.add_picture(img_stream, width=Inches(5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
        
        doc.add_paragraph()
        branding = doc.add_paragraph('Prepared by: Learning and Behaviour Unit')
        branding.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in branding.runs:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(14, 165, 233)
        
        doc.add_page_break()
        
        # SCHOOL NAME HEADER
        school_name = doc.add_heading('Cowandilla Learning Centre', 0)
        school_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in school_name.runs:
            run.font.color.rgb = GREEN_RGB
            run.font.size = Pt(18)
            set_arial(run)
        
        doc.add_paragraph()
        
        # STUDENT INFORMATION
        heading = doc.add_heading('Student Information', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_table.rows[0].cells[0].text = 'Student:'
        info_table.rows[0].cells[1].text = student['name']
        info_table.rows[1].cells[0].text = 'Program:'
        info_table.rows[1].cells[1].text = PROGRAM_NAMES.get(student['program'], student['program'])
        info_table.rows[2].cells[0].text = 'Grade:'
        info_table.rows[2].cells[1].text = student['grade']
        info_table.rows[3].cells[0].text = 'Analysis completed on:'
        info_table.rows[3].cells[1].text = datetime.now().strftime('%d %B %Y')
        info_table.rows[4].cells[0].text = 'Data Period:'
        info_table.rows[4].cells[1].text = f"{full_df['date_parsed'].min().strftime('%d/%m/%Y')} - {full_df['date_parsed'].max().strftime('%d/%m/%Y')}"
        
        doc.add_paragraph()
        
        # SUMMARY
        heading = doc.add_heading('Summary', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        summary = doc.add_paragraph()
        summary.add_run('Total Incidents: ').bold = True
        summary.add_run(f"{len(full_df)}\n")
        summary.add_run('Critical Incidents: ').bold = True
        summary.add_run(f"{len(full_df[full_df['incident_type'] == 'Critical'])}\n")
        summary.add_run('Average Severity: ').bold = True
        summary.add_run(f"{full_df['severity'].mean():.2f}/5\n")
        summary.add_run('Risk Level: ').bold = True
        summary.add_run(f"{risk_level} ({risk_score}/100)")
        
        doc.add_paragraph()
        
        # KEY FINDINGS
        heading = doc.add_heading('Key Findings', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        findings = doc.add_paragraph()
        findings.add_run('Behaviours of Concern: ').bold = True
        findings.add_run(f"{top_beh}\n\n")
        findings.add_run('Most Common Trigger: ').bold = True
        findings.add_run(f"{top_ant}\n\n")
        findings.add_run('Hotspot Location: ').bold = True
        findings.add_run(f"{top_loc}\n\n")
        findings.add_run('Occurs mainly in the: ').bold = True
        findings.add_run(f"{top_session}")
        
        doc.add_page_break()
        
        # VISUAL ANALYTICS
        heading = doc.add_heading('Visual Analytics', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        doc.add_paragraph('The following graphs provide visual representation of incident patterns and trends.')
        
        plt.style.use('default')
        from matplotlib.ticker import MaxNLocator
        
        # GRAPH 1: DAILY INCIDENT FREQUENCY (BAR CHART)
        heading = doc.add_heading('1. Daily Incident Frequency', 2)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        daily = full_df.groupby(full_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig, ax = plt.subplots(figsize=(10, 4), dpi=150)
        ax.bar(daily["date_parsed"], daily["count"], color='#334155', width=0.8, edgecolor='white', linewidth=0.5)
        ax.set_xlabel('Date', fontsize=11, fontweight='bold')
        ax.set_ylabel('Incident Count', fontsize=11, fontweight='bold')
        ax.grid(True, alpha=0.3, linestyle='--', axis='y')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        doc.add_paragraph("Daily incident frequency shows when behaviours occur most often.")
        doc.add_paragraph()
        
        # GRAPH 2: MOST COMMON BEHAVIOURS
        heading = doc.add_heading('2. Most Common Behaviours', 2)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        beh_counts = full_df["behaviour_type"].value_counts().head(5)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        ax.barh(beh_counts.index, beh_counts.values, color='#334155')
        ax.set_xlabel('Incident Count', fontsize=11, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.xaxis.set_major_locator(MaxNLocator(integer=True))
        for i, v in enumerate(beh_counts.values):
            ax.text(v + 0.3, i, str(int(v)), va='center', fontweight='bold')
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        doc.add_paragraph(f"Primary behaviour of concern: {beh_counts.index[0]} ({int(beh_counts.values[0])} incidents).")
        doc.add_paragraph()
        
        # GRAPH 3: MOST COMMON TRIGGERS
        heading = doc.add_heading('3. Most Common Triggers', 2)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        ant_counts = full_df["antecedent"].value_counts().head(5)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        ax.barh(ant_counts.index, ant_counts.values, color='#475569')
        ax.set_xlabel('Incident Count', fontsize=11, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.xaxis.set_major_locator(MaxNLocator(integer=True))
        for i, v in enumerate(ant_counts.values):
            ax.text(v + 0.3, i, str(int(v)), va='center', fontweight='bold')
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        doc.add_paragraph(f"Key trigger: {ant_counts.index[0]}.")
        doc.add_paragraph()
        
        # GRAPH 4: SEVERITY OVER TIME (COLOR-CODED, NO TREND LINE)
        heading = doc.add_heading('4. Severity Over Time', 2)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        fig, ax = plt.subplots(figsize=(10, 4), dpi=150)
        colors = {1: '#10b981', 2: '#3b82f6', 3: '#f59e0b', 4: '#ef4444', 5: '#7f1d1d'}
        for sev_level in [1, 2, 3, 4, 5]:
            sev_data = full_df[full_df['severity'] == sev_level]
            if len(sev_data) > 0:
                ax.scatter(sev_data["date_parsed"], sev_data["severity"], 
                          alpha=0.7, s=80, color=colors[sev_level], 
                          label=f'Level {sev_level}', edgecolors='white', linewidth=0.5)
        ax.set_xlabel('Date', fontsize=11, fontweight='bold')
        ax.set_ylabel('Severity', fontsize=11, fontweight='bold')
        ax.set_ylim(0, 5.5)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.legend(loc='upper right', frameon=False)
        ax.grid(True, alpha=0.3, linestyle='--', axis='y')
        plt.xticks(rotation=45, ha='right')
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        doc.add_paragraph("Severity levels shown by colour: Green (Level 1-2 = Minor), Blue (Level 3 = Moderate), Orange (Level 4 = Serious), Red (Level 5 = Critical).")
        doc.add_paragraph()
        
        # GRAPH 5: LOCATION HOTSPOTS
        heading = doc.add_heading('5. Location Hotspots', 2)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        loc_counts = full_df["location"].value_counts().head(5)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        ax.barh(loc_counts.index, loc_counts.values, color='#64748b')
        ax.set_xlabel('Incident Count', fontsize=11, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.xaxis.set_major_locator(MaxNLocator(integer=True))
        for i, v in enumerate(loc_counts.values):
            ax.text(v + 0.3, i, str(int(v)), va='center', fontweight='bold')
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        doc.add_paragraph(f"Most incidents occur in: {loc_counts.index[0]}.")
        doc.add_paragraph()
        
        # GRAPH 6: TIME OF DAY PATTERNS (IMPROVED)
        heading = doc.add_heading('6. Time of Day Patterns', 2)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
            
        session_counts = full_df["session"].value_counts()
        session_order = ['Morning', 'Middle', 'Afternoon']
        session_counts = session_counts.reindex(session_order, fill_value=0)
        fig, ax = plt.subplots(figsize=(8, 4), dpi=150)
        bars = ax.bar(session_counts.index, session_counts.values, color='#475569', edgecolor='white', linewidth=1.5)
        ax.set_ylabel('Incident Count', fontsize=11, fontweight='bold')
        ax.set_xlabel('Time of Day', fontsize=11, fontweight='bold')
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        for i, (bar, v) in enumerate(zip(bars, session_counts.values)):
            height = bar.get_height()
            ax.text(bar.get_x() + bar.get_width()/2., height + 0.3,
                   str(int(v)), ha='center', va='bottom', fontweight='bold', fontsize=11)
        plt.tight_layout()
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
        img_buffer.seek(0)
        doc.add_picture(img_buffer, width=Inches(6))
        plt.close()
        peak_session = session_counts.idxmax()
        peak_count = int(session_counts.max())
        doc.add_paragraph(f"Peak time: {peak_session} session with {peak_count} incidents. This pattern helps identify when additional support is most needed.")
        
        doc.add_page_break()
        
        # CLINICAL INTERPRETATION (ENHANCED)
        heading = doc.add_heading('Clinical Interpretation', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        intro = doc.add_paragraph()
        intro.add_run('This analysis is grounded in evidence-based frameworks that help us understand and support student behaviour. The following interpretation uses:')
        doc.add_paragraph('‚Ä¢ Applied Behaviour Analysis (ABA) - understanding what triggers and maintains behaviours', style='List Bullet')
        doc.add_paragraph('‚Ä¢ Trauma-Informed Practice - recognizing that behaviour is communication and often a response to stress', style='List Bullet')
        doc.add_paragraph('‚Ä¢ Berry Street Education Model - a whole-school approach to wellbeing and engagement', style='List Bullet')
        doc.add_paragraph('‚Ä¢ Crisis Prevention Institute (CPI) principles - de-escalation and maintaining dignity', style='List Bullet')
        doc.add_paragraph()
        
        # PATTERN ANALYSIS
        pattern_heading = doc.add_heading('Understanding the Patterns', 2)
        for run in pattern_heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        pattern = doc.add_paragraph()
        pattern.add_run('What the data tells us:\n').bold = True
        total_incidents = len(full_df)
        morning_pct = (len(full_df[full_df['session'] == 'Morning']) / total_incidents * 100) if total_incidents > 0 else 0
        middle_pct = (len(full_df[full_df['session'] == 'Middle']) / total_incidents * 100) if total_incidents > 0 else 0
        afternoon_pct = (len(full_df[full_df['session'] == 'Afternoon']) / total_incidents * 100) if total_incidents > 0 else 0
        pattern_text = pattern.add_run(
            f"Based on analysis of {total_incidents} recorded incidents, {student['name']} experiences the most difficulty "
            f"when '{top_ant}' occurs. This happens most frequently in {top_loc}, particularly during the {top_session} session. "
            f"\n\nTime of day breakdown shows: Morning ({morning_pct:.0f}% of incidents), Middle of day ({middle_pct:.0f}%), "
            f"and Afternoon ({afternoon_pct:.0f}%). This pattern suggests that {student['name']}'s ability to regulate and cope "
            f"is affected by time of day, likely due to factors like fatigue, hunger, sensory overload, or accumulated stress."
            f"\n\nThe behaviour '{top_beh}' is the primary concern. From a trauma-informed perspective, this behaviour is "
            f"{student['name']}'s way of communicating an unmet need or responding to feeling unsafe or overwhelmed. "
            f"It is not 'naughtiness' or 'choosing' to misbehave - it is a stress response."
        )
        set_arial(pattern_text)
        doc.add_paragraph()
        
        # BERRY STREET EDUCATION MODEL
        berry_heading = doc.add_heading('Berry Street Education Model Framework', 2)
        for run in berry_heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        berry_intro = doc.add_paragraph()
        berry_intro_text = berry_intro.add_run(
            "The Berry Street Education Model is an evidence-based whole-school approach developed in partnership "
            "with Melbourne University. It recognizes that students cannot learn effectively when they are stressed, "
            "unsafe, or disconnected. The model provides a sequential framework across five domains that must be "
            "addressed in order:\n"
        )
        set_arial(berry_intro_text)
        
        # Domain 1: BODY
        domain1 = doc.add_paragraph()
        d1_title = domain1.add_run('\n1. BODY Domain (Physical and Emotional Regulation):\n')
        d1_title.bold = True
        set_arial(d1_title)
        d1_text = domain1.add_run(
            f"This is the foundation. {student['name']} must feel physically and emotionally safe and regulated before "
            "anything else can happen. This means:\n"
            "‚Ä¢ Helping the student recognize their own body signals (heart racing, muscles tight, breathing fast)\n"
            "‚Ä¢ Teaching and practicing calming strategies (deep breathing, movement breaks, sensory tools)\n"
            "‚Ä¢ Creating predictable routines so the nervous system feels safe\n"
            "‚Ä¢ Providing regulation breaks BEFORE dysregulation occurs\n"
            f"‚Ä¢ Recognizing that during {top_session}, {student['name']} may need extra body-based support\n\n"
            "WHY THIS MATTERS: When stressed, the brain's 'thinking centre' goes offline and the 'survival centre' takes over. "
            "A dysregulated student cannot access learning, problem-solving, or relationship skills. "
            "We must help them get regulated first."
        )
        set_arial(d1_text)
        
        # Domain 2: RELATIONSHIP
        domain2 = doc.add_paragraph()
        d2_title = domain2.add_run('\n2. RELATIONSHIP Domain (Connection and Trust):\n')
        d2_title.bold = True
        set_arial(d2_title)
        d2_text = domain2.add_run(
            "Once physically regulated, students need safe, predictable relationships. This means:\n"
            f"‚Ä¢ One key adult who {student['name']} can trust and turn to when struggling\n"
            "‚Ä¢ Consistent, calm responses even during difficult behaviour\n"
            "‚Ä¢ Seeing behaviour as communication, not defiance\n"
            "‚Ä¢ Maintaining connection even when setting boundaries\n"
            "‚Ä¢ Understanding that damaged relationships must be repaired before learning can resume\n\n"
            "WHY THIS MATTERS: Behaviour often escalates when students feel disconnected, misunderstood, or unsafe "
            "in relationships. A strong relationship with at least one trusted adult is protective and helps students "
            "regulate their emotions and behaviour."
        )
        set_arial(d2_text)
        
        # Domain 3: STAMINA
        domain3 = doc.add_paragraph()
        d3_title = domain3.add_run('\n3. STAMINA Domain (Persistence and Resilience):\n')
        d3_title.bold = True
        set_arial(d3_title)
        d3_text = domain3.add_run(
            f"With regulation and connection in place, we can build {student['name']}'s capacity to persist with challenges:\n"
            "‚Ä¢ Breaking tasks into smaller, achievable steps\n"
            "‚Ä¢ Celebrating effort, not just outcomes\n"
            "‚Ä¢ Teaching that mistakes are part of learning\n"
            "‚Ä¢ Building confidence through success experiences\n"
            "‚Ä¢ Gradually increasing expectations as capacity grows\n\n"
            "WHY THIS MATTERS: Students who have experienced trauma or chronic stress often have learned that "
            "'trying' leads to failure or shame. We must rebuild their belief that effort matters and that they are capable."
        )
        set_arial(d3_text)
        
        # Domain 4: ENGAGEMENT
        domain4 = doc.add_paragraph()
        d4_title = domain4.add_run('\n4. ENGAGEMENT Domain (Active Learning Participation):\n')
        d4_title.bold = True
        set_arial(d4_title)
        d4_text = domain4.add_run(
            "With the first three domains secure, students can engage meaningfully in learning:\n"
            "‚Ä¢ Making learning relevant and purposeful\n"
            "‚Ä¢ Providing choice and autonomy\n"
            "‚Ä¢ Using strengths and interests\n"
            "‚Ä¢ Creating positive relationships with learning\n"
            f"‚Ä¢ Recognizing that during {top_session}, engagement may need additional scaffolding\n\n"
            "WHY THIS MATTERS: Students cannot engage in learning when dysregulated, disconnected, or defeated. "
            "Engagement is a result of getting the foundation domains right, not something we can demand."
        )
        set_arial(d4_text)
        
        # Domain 5: CHARACTER
        domain5 = doc.add_paragraph()
        d5_title = domain5.add_run('\n5. CHARACTER Domain (Values and Contribution):\n')
        d5_title.bold = True
        set_arial(d5_title)
        d5_text = domain5.add_run(
            "The final domain focuses on purpose and positive contribution:\n"
            "‚Ä¢ Developing empathy and perspective-taking\n"
            "‚Ä¢ Understanding impact of actions on others\n"
            "‚Ä¢ Finding ways to contribute positively\n"
            "‚Ä¢ Building identity as a capable, valued person\n\n"
            "WHY THIS MATTERS: This is NOT about 'being good' or compliance. It's about helping students "
            "develop a positive sense of self and their place in the community. This domain only works when "
            "the foundation domains are solid."
        )
        set_arial(d5_text)
        doc.add_paragraph()
        
        berry_application = doc.add_paragraph()
        ba_title = berry_application.add_run(f'Application to {student["name"]}:\n')
        ba_title.bold = True
        set_arial(ba_title)
        ba_text = berry_application.add_run(
            f"Currently, our focus must be on BODY and RELATIONSHIP domains. The data shows {student['name']} "
            f"is dysregulated during {top_session}, particularly when '{top_ant}' occurs. "
            "We cannot expect engagement or character development until regulation and connection are secure. "
            "All recommendations in this plan prioritize these foundation domains first."
        )
        set_arial(ba_text)
        doc.add_paragraph()
        
        # CPI PRINCIPLES
        cpi_heading = doc.add_heading('Crisis Prevention Institute (CPI) Principles', 2)
        for run in cpi_heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        cpi_intro = doc.add_paragraph()
        cpi_intro_text = cpi_intro.add_run(
            "CPI provides evidence-based training in de-escalation and crisis prevention. The core principles guide "
            "how we respond when students are escalating or in crisis:\n"
        )
        set_arial(cpi_intro_text)
        
        # CPI Principle 1
        cpi1 = doc.add_paragraph()
        cpi1_title = cpi1.add_run('\nBehaviour is Communication:\n')
        cpi1_title.bold = True
        set_arial(cpi1_title)
        cpi1_text = cpi1.add_run(
            f"When {student['name']} displays '{top_beh}', they are communicating something important. "
            "They might be saying: 'I'm overwhelmed,' 'I feel unsafe,' 'I don't understand,' 'I need help,' "
            "or 'I'm hungry/tired/stressed.' Our job is to understand the message, not just stop the behaviour."
        )
        set_arial(cpi1_text)
        
        # CPI Principle 2
        cpi2 = doc.add_paragraph()
        cpi2_title = cpi2.add_run('\nSupportive Stance:\n')
        cpi2_title.bold = True
        set_arial(cpi2_title)
        cpi2_text = cpi2.add_run(
            "How we position our body matters. Stand at an angle (not directly facing), give space (don't crowd), "
            "keep hands visible and open, stay at or below eye level. Use a low, slow, calm voice. "
            "Your body language should say: 'I'm here to help, you are safe with me.'"
        )
        set_arial(cpi2_text)
        
        # CPI Principle 3
        cpi3 = doc.add_paragraph()
        cpi3_title = cpi3.add_run('\nMaintain Dignity:\n')
        cpi3_title.bold = True
        set_arial(cpi3_title)
        cpi3_text = cpi3.add_run(
            "Never shame, embarrass, or humiliate. Don't have an audience for correction. "
            "Separate the behaviour from the person. The message should always be: 'I care about you, "
            "even when your behaviour is difficult.'"
        )
        set_arial(cpi3_text)
        
        # CPI Principle 4
        cpi4 = doc.add_paragraph()
        cpi4_title = cpi4.add_run('\nEarly Intervention:\n')
        cpi4_title.bold = True
        set_arial(cpi4_title)
        cpi4_text = cpi4.add_run(
            f"The data shows that '{top_ant}' is a key trigger. CPI teaches us to intervene early - "
            "at the first signs of escalation, before crisis. This might mean: offering a break, changing the task, "
            "providing reassurance, or simply moving to a quieter space. Early intervention prevents crisis."
        )
        set_arial(cpi4_text)
        
        # CPI Principle 5
        cpi5 = doc.add_paragraph()
        cpi5_title = cpi5.add_run('\nCo-Regulation:\n')
        cpi5_title.bold = True
        set_arial(cpi5_title)
        cpi5_text = cpi5.add_run(
            f"Students like {student['name']} often cannot regulate themselves when dysregulated. "
            "They need an adult to stay calm and lend them their regulation. Your calm becomes their calm. "
            "This is why adult self-regulation is essential - you cannot co-regulate if you are also dysregulated."
        )
        set_arial(cpi5_text)
        doc.add_paragraph()
        
        cpi_application = doc.add_paragraph()
        ca_title = cpi_application.add_run(f'Application to {student["name"]}:\n')
        ca_title.bold = True
        set_arial(ca_title)
        ca_text = cpi_application.add_run(
            f"When incidents occur in {top_loc} during {top_session}, staff should use CPI principles: "
            "approach calmly with supportive stance, offer choices to maintain dignity, intervene early "
            f"when '{top_ant}' is present, and provide co-regulation rather than consequences. "
            "The goal is always to help the student return to their window of tolerance, not to punish."
        )
        set_arial(ca_text)
        
        doc.add_page_break()
        
        # ================================================================
        # AUSTRALIAN CURRICULUM CAPABILITIES SECTION
        # ================================================================
        
        heading = doc.add_heading('Australian Curriculum - General Capabilities', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        ac_intro = doc.add_paragraph()
        ac_intro_run = ac_intro.add_run(
            "The Australian Curriculum includes General Capabilities that are developed across all learning areas. "
            "Behaviour support is directly connected to the Personal and Social Capability, which encompasses "
            "self-awareness, self-management, social awareness, and social management."
        )
        set_arial(ac_intro_run)
        
        doc.add_paragraph()
        
        # Get AC information for this student's primary behaviour
        ac_info = get_ac_capability_for_behaviour(top_beh) if 'get_ac_capability_for_behaviour' in dir() else None
        
        if ac_info:
            # Primary Capability
            cap_heading = doc.add_heading('Primary Capability Focus', 2)
            for run in cap_heading.runs:
                run.font.color.rgb = GREEN_RGB
                set_arial(run)
            
            cap_para = doc.add_paragraph()
            cap_title = cap_para.add_run(f"{ac_info['name']} ({ac_info['code']})")
            cap_title.bold = True
            cap_title.font.size = Pt(12)
            
            elements_para = doc.add_paragraph()
            elements_para.add_run(f"Focus Elements: {', '.join(ac_info['elements'])}")
            
            doc.add_paragraph()
            
            # Skills to Develop
            skills_heading = doc.add_heading('Skills to Develop (AC-Aligned)', 2)
            for run in skills_heading.runs:
                run.font.color.rgb = GREEN_RGB
                set_arial(run)
            
            for skill in ac_info['skills_to_develop']:
                doc.add_paragraph(f"‚Ä¢ {skill}", style='List Bullet')
            
            doc.add_paragraph()
            
            # AC Descriptors
            desc_heading = doc.add_heading('Curriculum Descriptors', 2)
            for run in desc_heading.runs:
                run.font.color.rgb = GREEN_RGB
                set_arial(run)
            
            for desc in ac_info['ac_descriptors']:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(desc)
                run.italic = True
        
        else:
            # Default PSC content
            psc_heading = doc.add_heading('Personal and Social Capability', 2)
            for run in psc_heading.runs:
                run.font.color.rgb = GREEN_RGB
                set_arial(run)
            
            psc_elements = doc.add_paragraph()
            psc_elements.add_run("Key Elements for Behaviour Support:").bold = True
            
            doc.add_paragraph("‚Ä¢ Self-awareness: Recognise emotions, understand themselves as learners", style='List Bullet')
            doc.add_paragraph("‚Ä¢ Self-management: Express emotions appropriately, develop self-discipline", style='List Bullet')
            doc.add_paragraph("‚Ä¢ Social awareness: Appreciate diverse perspectives, understand relationships", style='List Bullet')
            doc.add_paragraph("‚Ä¢ Social management: Communicate effectively, resolve conflict", style='List Bullet')
        
        doc.add_paragraph()
        
        # Learning Goals
        goals_heading = doc.add_heading('AC-Aligned Learning Goals', 2)
        for run in goals_heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        goals_intro = doc.add_paragraph()
        goals_intro_run = goals_intro.add_run(f"Suggested goals for {student['name']} based on behaviour patterns and grade level ({student['grade']}):")
        goals_intro_run.italic = True
        set_arial(goals_intro_run)
        
        # Generate grade-appropriate goals
        learning_goals = generate_ac_learning_goals(top_beh, top_ant, student['grade']) if 'generate_ac_learning_goals' in dir() else []
        
        if learning_goals:
            for goal in learning_goals:
                p = doc.add_paragraph(style='List Bullet')
                skill_run = p.add_run(f"{goal['skill']}: ")
                skill_run.bold = True
                p.add_run(f"{goal['goal']} ")
                desc_run = p.add_run(f"({goal['ac_descriptor']})")
                desc_run.italic = True
        else:
            doc.add_paragraph("‚Ä¢ Emotional recognition: Student will identify and name emotions when prompted", style='List Bullet')
            doc.add_paragraph("‚Ä¢ Help-seeking: Student will use a help-seeking strategy independently", style='List Bullet')
            doc.add_paragraph("‚Ä¢ Self-regulation: Student will use a calming strategy when dysregulated", style='List Bullet')
        
        doc.add_paragraph()
        
        # Connection note
        note_para = doc.add_paragraph()
        note_title = note_para.add_run("Important: ")
        note_title.bold = True
        note_para.add_run(
            "Progress in Personal and Social Capability should be documented alongside behaviour data. "
            "This demonstrates growth in underlying skills, not just reduction in incidents."
        )
        
        doc.add_page_break()
        
        # ================================================================
        # EVIDENCE-BASED RECOMMENDATIONS
        # ================================================================
        
        heading = doc.add_heading('Evidence-Based Recommendations', 1)
        for run in heading.runs:
            run.font.color.rgb = GREEN_RGB
            set_arial(run)
        
        heading = doc.add_heading('1. Body Domain (Regulation)', 2)

        for run in heading.runs:

            run.font.color.rgb = GREEN_RGB

            set_arial(run)
        doc.add_paragraph(f"‚Ä¢ Regulated start before {top_session}", style='List Bullet')
        doc.add_paragraph(f"‚Ä¢ Visual check-in before '{top_ant}'", style='List Bullet')
        doc.add_paragraph(f"‚Ä¢ Environmental modification in {top_loc}", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Sensory regulation opportunities", style='List Bullet')
        
        heading = doc.add_heading('2. Relationship Domain (Connection)', 2)

        
        for run in heading.runs:

        
            run.font.color.rgb = GREEN_RGB

        
            set_arial(run)
        doc.add_paragraph("‚Ä¢ Supportive Stance with low, slow voice", style='List Bullet')
        doc.add_paragraph("‚Ä¢ One key adult maintains connection", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Acknowledge feelings", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Offer choices for control", style='List Bullet')
        
        heading = doc.add_heading('3. Stamina Domain (Persistence)', 2)

        
        for run in heading.runs:

        
            run.font.color.rgb = GREEN_RGB

        
            set_arial(run)
        doc.add_paragraph("‚Ä¢ Teach help-seeking with visual cues", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Practice requesting breaks", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Emotional literacy skills", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Build coping strategies", style='List Bullet')
        
        heading = doc.add_heading('4. SMART Goal (AC-Aligned)', 2)

        
        for run in heading.runs:

        
            run.font.color.rgb = GREEN_RGB

        
            set_arial(run)
        goal = doc.add_paragraph()
        goal.add_run('Measurable: ').bold = True
        goal.add_run(f"Over 5 weeks, {student['name']} will use a help-seeking strategy in 4/5 opportunities with support.")
        doc.add_paragraph()
        goal_ac = doc.add_paragraph()
        goal_ac_run = goal_ac.add_run("AC Alignment: PSC - Self-management: 'Work independently and show initiative'")
        goal_ac_run.italic = True
        doc.add_paragraph()
        doc.add_paragraph('Review Date: ' + (datetime.now() + timedelta(weeks=5)).strftime('%d %B %Y'))
        
        doc.add_page_break()
        
        # FOOTER
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('\n\nLearning and Behaviour Unit\n')
        footer_run.font.size = Pt(10)
        footer_run.font.bold = True
        footer_run.font.color.rgb = RGBColor(14, 165, 233)
        
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer2_run = footer2.add_run('Evidence-based: ABA, Trauma-Informed, Berry Street, CPI, Australian Curriculum\n')
        footer2_run.font.size = Pt(9)
        footer2_run.font.color.rgb = RGBColor(100, 116, 139)
        
        footer3 = doc.add_paragraph()
        footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer3_run = footer3.add_run(datetime.now().strftime('%d %B %Y'))
        footer3_run.font.size = Pt(9)
        footer3_run.font.color.rgb = RGBColor(100, 116, 139)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
        
    except Exception as e:
        import traceback
        st.error(f"BAP Error: {e}")
        st.error(traceback.format_exc())
        return None


# ============================================
# SUPABASE DATABASE FUNCTIONS
# ============================================

def hash_password(plain_password: str) -> str:
    """Hash a password using bcrypt"""
    salt = bcrypt.gensalt(rounds=12)
    hashed = bcrypt.hashpw(plain_password.encode('utf-8'), salt)
    return hashed.decode('utf-8')

# FIXED load_students_from_db FUNCTION
# Replace the existing load_students_from_db function (around line 1060-1085)

def load_students_from_db():
    """Load students from Supabase database"""
    if not supabase:
        return MOCK_STUDENTS  # Fallback to mock data
    
    try:
        response = supabase.table('students').select('*').execute()
        students = []
        for row in response.data:
            # Convert grade from integer back to string format
            grade_num = row['grade']
            if grade_num == 0:
                grade_str = 'R'
            else:
                grade_str = f'Y{grade_num}'
            
            # Handle first_name/last_name - construct name if not present
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            full_name = row.get('name', '')
            
            # If first_name/last_name not in DB, parse from name
            if not first_name and full_name:
                parts = full_name.split()
                first_name = parts[0] if parts else ''
                last_name = ' '.join(parts[1:]) if len(parts) > 1 else ''
            
            # If name not in DB, construct from first/last
            if not full_name and (first_name or last_name):
                full_name = f"{first_name} {last_name}".strip()
            
            students.append({
                "id": str(row['id']),
                "first_name": first_name,
                "last_name": last_name,
                "name": full_name,
                "edid": row['edid'],
                "grade": grade_str,  # Convert back to Y1, Y2, R format
                "dob": row['dob'],
                "program": row['program'],
                "placement_start": row['placement_start'],
                "placement_end": row['placement_end']
            })
        return students
    except Exception as e:
        st.error(f"Error loading students: {e}")
        return []

def save_student_to_db(student):
    """Save a student to Supabase database"""
    if not supabase:
        return False
    
    try:
        # Convert grade to just the number if it starts with Y
        grade_value = student['grade']
        if isinstance(grade_value, str):
            if grade_value.startswith('Y'):
                grade_value = grade_value[1:]
            elif grade_value == 'R':
                grade_value = 0
        
        data = {
            "first_name": student.get('first_name', student['name'].split()[0] if student['name'] else ''),
            "last_name": student.get('last_name', ' '.join(student['name'].split()[1:]) if len(student['name'].split()) > 1 else ''),
            "name": student['name'],
            "edid": student['edid'],
            "grade": int(grade_value) if str(grade_value).isdigit() else 0,
            "year_level": int(grade_value) if str(grade_value).isdigit() else 0,
            "dob": student['dob'],
            "program": student['program'],
            "placement_start": student['placement_start'],
            "placement_end": student.get('placement_end')
        }
        
        if 'id' in student and student['id'].startswith('stu_'):
            supabase.table('students').insert(data).execute()
        else:
            supabase.table('students').update(data).eq('id', student['id']).execute()
        
        return True
        
    except Exception as e:
        st.error(f"Error saving student: {e}")
        return False
def delete_student_from_db(student_id):
    """Delete a student from Supabase database"""
    if not supabase:
        return False
    
    try:
        supabase.table('students').delete().eq('id', student_id).execute()
        return True
    except Exception as e:
        st.error(f"Error deleting student: {e}")
        return False

def load_staff_from_db():
    """Load staff from Supabase database"""
    if not supabase:
        return MOCK_STAFF
    
    try:
        response = supabase.table('staff').select('*').execute()
        staff = []
        for row in response.data:
            # Handle first_name/last_name - construct name if not present
            first_name = row.get('first_name', '')
            last_name = row.get('last_name', '')
            full_name = row.get('name', '')
            
            # If first_name/last_name not in DB, parse from name
            if not first_name and full_name:
                parts = full_name.split()
                first_name = parts[0] if parts else ''
                last_name = ' '.join(parts[1:]) if len(parts) > 1 else ''
            
            # If name not in DB, construct from first/last
            if not full_name and (first_name or last_name):
                full_name = f"{first_name} {last_name}".strip()
            
            staff.append({
                "id": str(row['id']),
                "first_name": first_name,
                "last_name": last_name,
                "name": full_name,
                "email": row['email'],
                "password": row.get('password'),  # Keep for backward compatibility
                "password_hash": row.get('password_hash', row.get('password')),  # Use password_hash if available
                "role": row['role'],
                "program": row.get('program'),
                "phone": row.get('phone'),
                "notes": row.get('notes'),
                "receive_critical_emails": row.get('receive_critical_emails', True),
                "created_date": row.get('created_at', '')[:10] if row.get('created_at') else None
            })
        return staff  # Return what we have from database
    except Exception as e:
        st.error(f"Error loading staff: {e}")
        return []  # Return empty list on error

def save_staff_to_db(staff_member):
    """Save a staff member to Supabase database"""
    if not supabase:
        return False
    
    try:
        data = {
            "first_name": staff_member.get('first_name', staff_member['name'].split()[0] if staff_member['name'] else ''),
            "last_name": staff_member.get('last_name', ' '.join(staff_member['name'].split()[1:]) if len(staff_member['name'].split()) > 1 else ''),
            "name": staff_member['name'],
            "email": staff_member['email'],
            "password": staff_member['password'],
            "role": staff_member['role'],
            "program": staff_member.get('program'),
            "phone": staff_member.get('phone'),
            "notes": staff_member.get('notes'),
            "receive_critical_emails": staff_member.get('receive_critical_emails', True)
        }
        
        if 'id' in staff_member and staff_member['id'].startswith('staff_'):
            # New staff (generated ID from app)
            supabase.table('staff').insert(data).execute()
        else:
            # Existing staff (UUID from database)
            supabase.table('staff').update(data).eq('id', staff_member['id']).execute()
        return True
    except Exception as e:
        st.error(f"Error saving staff: {e}")
        return False

def delete_staff_from_db(staff_id):
    """Delete a staff member from Supabase database"""
    if not supabase:
        return False
    
    try:
        supabase.table('staff').delete().eq('id', staff_id).execute()
        return True
    except Exception as e:
        st.error(f"Error deleting staff: {e}")
        return False

def load_incidents_from_db(student_id=None):
    """Load incidents from Supabase database"""
    if not supabase:
        return []
    
    try:
        query = supabase.table('incidents').select('*')
        if student_id:
            query = query.eq('student_id', student_id)
        response = query.execute()
        
        incidents = []
        for row in response.data:
            incidents.append({
                "id": str(row['id']),
                "student_id": str(row['student_id']),
                "date": row['date'],
                "time": row['time'],
                "day": row['day_of_week'],
                "session": row['session'],
                "location": row['location'],
                "behaviour_type": row['behaviour_type'],
                "antecedent": row['antecedent'],
                "intervention": row['intervention'],  # Already array in DB
                "severity": row['severity'],
                "reported_by": str(row['reported_by']) if row.get('reported_by') else None,
                "description": row.get('description', ''),
                "duration_minutes": row.get('duration_minutes'),
                "hypothesis_function": row.get('hypothesis_function'),
                "hypothesis_item": row.get('hypothesis_item'),
                "is_critical": row.get('is_critical', False)
            })
        return incidents
    except Exception as e:
        st.error(f"Error loading incidents: {e}")
        return []

def save_incident_to_db(incident):
    """Save an incident to Supabase database"""
    if not supabase:
        return False
    
    try:
        data = {
            "student_id": incident['student_id'],
            "date": incident['date'],
            "time": incident['time'],
            "day_of_week": incident['day'],
            "session": incident['session'],
            "location": incident['location'],
            "behaviour_type": incident['behaviour_type'],
            "antecedent": incident['antecedent'],
            "intervention": incident['intervention'],
            "severity": incident['severity'],
            "reported_by": incident.get('reported_by'),
            "description": incident.get('description', ''),
            "duration_minutes": incident.get('duration_minutes'),
            "hypothesis_function": incident.get('hypothesis_function'),
            "hypothesis_item": incident.get('hypothesis_item'),
            "is_critical": incident.get('is_critical', False)
        }
        
        if 'id' not in incident or not incident['id']:
            # New incident
            supabase.table('incidents').insert(data).execute()
        else:
            # Update existing
            supabase.table('incidents').update(data).eq('id', incident['id']).execute()
        return True
    except Exception as e:
        st.error(f"Error saving incident: {e}")
        return False

def load_critical_incidents_from_db(student_id=None):
    """Load critical incidents from Supabase database"""
    if not supabase:
        return []
    
    try:
        query = supabase.table('critical_incidents').select('*')
        if student_id:
            query = query.eq('student_id', student_id)
        response = query.execute()
        
        critical = []
        for row in response.data:
            critical.append({
                "id": str(row['id']),
                "student_id": str(row['student_id']),
                "severity": row['severity'],
                "reported_by": str(row['reported_by']) if row.get('reported_by') else None,
                "ABCH_primary": {
                    "location": row['primary_location'],
                    "context": row['primary_context'],
                    "time": row['primary_time'],
                    "behaviour": row['primary_behaviour'],
                    "consequence": row['primary_consequence'],
                    "hypothesis_function": row.get('primary_hypothesis_function'),
                    "hypothesis_item": row.get('primary_hypothesis_item')
                },
                "ABCH_additional": row.get('additional_entries', []),
                "outcomes": row['outcomes'],
                "sapol_reference": row.get('sapol_reference'),
                "admin_summary": row.get('admin_summary'),
                "created_at": row.get('created_at')
            })
        return critical
    except Exception as e:
        st.error(f"Error loading critical incidents: {e}")
        return []

def save_critical_incident_to_db(critical):
    """Save a critical incident to Supabase database"""
    if not supabase:
        return False
    
    try:
        primary = critical.get('ABCH_primary', {})
        data = {
            "student_id": critical['student_id'],
            "severity": critical['severity'],
            "reported_by": critical.get('reported_by'),
            "primary_location": primary.get('location', ''),
            "primary_context": primary.get('context', ''),
            "primary_time": primary.get('time', ''),
            "primary_behaviour": primary.get('behaviour', ''),
            "primary_consequence": primary.get('consequence', ''),
            "primary_hypothesis_function": primary.get('hypothesis_function'),
            "primary_hypothesis_item": primary.get('hypothesis_item'),
            "additional_entries": critical.get('ABCH_additional', []),
            "outcomes": critical.get('outcomes', []),
            "sapol_reference": critical.get('sapol_reference'),
            "admin_summary": critical.get('admin_summary')
        }
        
        if 'id' not in critical or not critical['id']:
            # New critical incident
            supabase.table('critical_incidents').insert(data).execute()
        else:
            # Update existing
            supabase.table('critical_incidents').update(data).eq('id', critical['id']).execute()
        return True
    except Exception as e:
        st.error(f"Error saving critical incident: {e}")
        return False


def init_state():
    ss = st.session_state
    if "logged_in" not in ss: ss.logged_in = False
    if "current_user" not in ss: ss.current_user = None
    if "current_page" not in ss: ss.current_page = "login"
    
    # Load from Supabase if available, otherwise use mock data
    if "students" not in ss: 
        ss.students = load_students_from_db() if supabase else MOCK_STUDENTS
    if "staff" not in ss: 
        ss.staff = load_staff_from_db() if supabase else MOCK_STAFF
    if "incidents" not in ss: 
        ss.incidents = load_incidents_from_db() if supabase else generate_mock_incidents(70)
    if "critical_incidents" not in ss: 
        ss.critical_incidents = load_critical_incidents_from_db() if supabase else []
    
    if "selected_program" not in ss: ss.selected_program = "JP"
    if "selected_student_id" not in ss: ss.selected_student_id = None
    if "current_incident_id" not in ss: ss.current_incident_id = None
    if "abch_rows" not in ss: ss.abch_rows = []
    if "show_critical_prompt" not in ss: ss.show_critical_prompt = False

def login_user(email: str, password: str) -> bool:
    """Login user with bcrypt password verification"""
    email = (email or "").strip().lower()
    password = (password or "").strip()
    
    st.write(f"üîç DEBUG: Attempting login with email: '{email}'")
    st.write(f"üîç DEBUG: Password length: {len(password)}")
    st.write(f"üîç DEBUG: Total staff in memory: {len(st.session_state.staff)}")
    
    if not email or not password:
        st.write("‚ùå DEBUG: Email or password empty")
        return False
    
    for idx, staff in enumerate(st.session_state.staff):
        staff_email = staff.get("email", "").lower()
        st.write(f"üîç DEBUG: Checking staff #{idx}: {staff_email}")
        
        if staff_email == email:
            st.write(f"‚úÖ DEBUG: Email match found!")
            st.write(f"üîç DEBUG: Staff has password field: {staff.get('password')}")
            st.write(f"üîç DEBUG: Staff has password_hash field: {staff.get('password_hash', '')[:30]}...")
            
            # Get the stored hash
            stored_hash = staff.get("password_hash", "")
            if not stored_hash:
                st.write("‚ö†Ô∏è DEBUG: No password_hash found, trying plain password")
                if staff.get("password") == password:
                    st.write("‚úÖ DEBUG: Plain password match!")
                    st.session_state.logged_in = True
                    st.session_state.current_user = staff
                    st.session_state.current_page = "landing"
                    return True
                else:
                    st.write(f"‚ùå DEBUG: Plain password mismatch. Expected: '{staff.get('password')}', Got: '{password}'")
                    continue
            
            # Verify password against bcrypt hash
            try:
                if isinstance(stored_hash, str):
                    stored_hash_bytes = stored_hash.encode('utf-8')
                else:
                    stored_hash_bytes = stored_hash
                    
                password_bytes = password.encode('utf-8')
                
                st.write(f"üîç DEBUG: Attempting bcrypt verification...")
                if bcrypt.checkpw(password_bytes, stored_hash_bytes):
                    st.write("‚úÖ DEBUG: Bcrypt verification SUCCESS!")
                    st.session_state.logged_in = True
                    st.session_state.current_user = staff
                    st.session_state.current_page = "landing"
                    return True
                else:
                    st.write("‚ùå DEBUG: Bcrypt verification FAILED")
                    if staff.get("password") == password:
                        st.write("‚úÖ DEBUG: Using plain password fallback - LOGIN SUCCESS!")
                        st.session_state.logged_in = True
                        st.session_state.current_user = staff
                        st.session_state.current_page = "landing"
                        return True
                    else:
                        st.write(f"‚ùå DEBUG: Plain password mismatch")
            except Exception as e:
                st.write(f"‚ö†Ô∏è DEBUG: Bcrypt error: {e}")
                if staff.get("password") == password:
                    st.write("‚úÖ DEBUG: Plain password fallback after exception - LOGIN SUCCESS!")
                    st.session_state.logged_in = True
                    st.session_state.current_user = staff
                    st.session_state.current_page = "landing"
                    return True
    
    st.write("‚ùå DEBUG: No matching staff found")
    return False
def go_to(page: str, **kwargs):
    if page not in VALID_PAGES: return
    st.session_state.current_page = page
    for k, v in kwargs.items():
        setattr(st.session_state, k, v)
    st.rerun()

def get_student(sid): 
    return next((s for s in st.session_state.students if s["id"] == sid), None)

def get_session_from_time(t): 
    return "Morning" if t.hour < 11 else "Middle" if t.hour < 13 else "Afternoon"

def generate_mock_incidents(n=70):
    incidents = []
    weights = {"stu_sy1": 12, "stu_py1": 10, "stu_sy2": 9, "stu_jp1": 8, "stu_py2": 7}
    pool = []
    for stu in MOCK_STUDENTS:
        pool.extend([stu] * weights.get(stu["id"], 5))
    for _ in range(n):
        stu = random.choice(pool)
        sev = random.choices([1, 2, 3, 4, 5], weights=[20, 35, 25, 15, 5])[0]
        dt = datetime.now() - timedelta(days=random.randint(0, 90))
        dt = dt.replace(hour=random.choices([9,10,11,12,13,14,15], weights=[10,15,12,8,12,18,10])[0], 
                       minute=random.randint(0,59), second=0)
        incidents.append({
            "id": str(uuid.uuid4()), "student_id": stu["id"], "student_name": stu["name"],
            "date": dt.date().isoformat(), "time": dt.time().strftime("%H:%M:%S"),
            "day": dt.strftime("%A"), "session": get_session_from_time(dt.time()),
            "location": random.choice(LOCATIONS), "behaviour_type": random.choice(BEHAVIOUR_TYPES),
            "antecedent": random.choice(ANTECEDENTS), 
            "intervention": [random.choice(INTERVENTIONS)],  # Changed to list
            "severity": sev, "reported_by": random.choice(MOCK_STAFF)["name"],
            "description": "Mock incident", "is_critical": sev >= 3, "duration_minutes": random.randint(2, 25)
        })
    return incidents

# PAGES
def render_login_page():
    st.markdown("## üîê Staff Login")

    email = st.text_input("Email Address", placeholder="your.email@example.com", key="login_email")
    password = st.text_input("Password", type="password", placeholder="Enter password", key="login_pass")
    if st.button("Login", type="primary", use_container_width=True):
        if login_user(email, password):
            st.success(f"Welcome {st.session_state.current_user['name']}!")
            st.rerun()
        else:
            st.error("Invalid credentials")

def render_landing_page():
    user = st.session_state.current_user or {}
    st.markdown(f"### üëã Welcome, {user.get('name', 'User')}")
    
    col1, col2 = st.columns([6, 1])
    with col2:
        if st.button("Logout", key="logout_btn"):
            st.session_state.logged_in = False
            st.session_state.current_page = "login"
            st.rerun()
    
    # Admin portal button for admin users
    if st.session_state.current_user.get("role") == "ADM":
        st.markdown("---")
        if st.button("üîß Admin Portal", use_container_width=True, key="goto_admin"):
            go_to("admin_portal")
    
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    with col1: st.metric("Students", len(st.session_state.students))
    with col2: st.metric("Total Incidents", len(st.session_state.incidents))
    with col3: st.metric("Critical", len([i for i in st.session_state.incidents if i.get("is_critical")]))
    
    st.markdown("---")
    st.markdown("### Select Program")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Junior Primary**")
        if st.button("Enter JP", use_container_width=True, type="primary", key="btn_jp"):
            go_to("program_students", selected_program="JP")
    with col2:
        st.markdown("**Primary Years**")
        if st.button("Enter PY", use_container_width=True, type="primary", key="btn_py"):
            go_to("program_students", selected_program="PY")
    with col3:
        st.markdown("**Senior Years**")
        if st.button("Enter SY", use_container_width=True, type="primary", key="btn_sy"):
            go_to("program_students", selected_program="SY")

def render_program_students_page():
    program = st.session_state.get("selected_program", "JP")
    st.markdown(f"## {PROGRAM_NAMES.get(program)} ‚Äî Students")
    
    col1, col2 = st.columns([6, 1])
    with col1:
        if st.button("‚¨Ö Back to Landing", key="back_students"):
            go_to("landing")
    
    students = [s for s in st.session_state.students if s["program"] == program]
    for stu in students:
        stu_incidents = [i for i in st.session_state.incidents if i["student_id"] == stu["id"]]
        with st.container(border=True):
            col1, col2, col3 = st.columns([3, 2, 2])
            with col1:
                st.markdown(f"**{stu['name']}**")
                st.caption(f"Grade {stu['grade']}")
            with col2:
                st.metric("Incidents", len(stu_incidents))
            with col3:
                if st.button("üìù Log", key=f"log_{stu['id']}", use_container_width=True):
                    go_to("incident_log", selected_student_id=stu["id"])
                if st.button("üìä Analysis", key=f"ana_{stu['id']}", use_container_width=True):
                    go_to("student_analysis", selected_student_id=stu["id"])

def render_incident_log_page():
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## üìù Incident Log ‚Äî {student['name']}")
    
    # Navigation buttons
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("‚¨Ö Back to Students", key="back_log_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("üè† Program Landing", key="home_log"):
            go_to("landing")
    
    show_severity_guide()
    
    # Check if critical form is required
    if st.session_state.show_critical_prompt:
        inc_info = st.session_state.get("last_incident_info", {})
        if inc_info.get("severity", 0) >= 3:
            st.error(f"‚ö†Ô∏è **Severity {inc_info['severity']} Detected** - Critical Incident ABCH Form Required")
        else:
            st.error("‚ö†Ô∏è **Critical Incident Flagged** - Critical Incident ABCH Form Required")
        st.info("Please complete the Critical Incident ABCH form to document this event fully.")
        
        # Only one button - REMOVED "Skip for Now"
        if st.button("üìã Complete Critical Form Now", type="primary", key="crit_now", use_container_width=True):
            st.session_state.show_critical_prompt = False
            go_to("critical_incident", current_incident_id=st.session_state.current_incident_id)
        st.markdown("---")
        st.stop()
    
    # INCIDENT FORM - Split to show hypothesis before severity
    st.markdown("### Log New Incident")
    
    # First section: Capture antecedent and behaviour for hypothesis generation
    col1, col2 = st.columns(2)
    with col1:
        behaviour_select = st.selectbox("Behaviour Type *", [""] + BEHAVIOUR_TYPES, key="inc_beh_select")
    with col2:
        antecedent_select = st.selectbox("Antecedent/Trigger *", [""] + ANTECEDENTS, key="inc_ant_select")
    
    # Display hypothesis if both are selected
    if behaviour_select and antecedent_select:
        hyp_ai = generate_hypothesis_ai(antecedent_select, behaviour_select, "")
        hypothesis_text = f"{hyp_ai['function']} {hyp_ai['item']}"
        st.info(f"üß† **Suggested Hypothesis:** {hypothesis_text}")
        # Store for form submission
        st.session_state.current_hypothesis = hyp_ai
    else:
        st.session_state.current_hypothesis = None
    
    # Main form with remaining fields
    with st.form("incident_form", clear_on_submit=True):
        col1, col2 = st.columns(2)
        with col1:
            inc_date = st.date_input("Date *", date.today(), key="inc_date", format="DD/MM/YYYY")
            inc_time = st.time_input("Time *", datetime.now().time(), key="inc_time")
            location = st.selectbox("Location *", [""] + LOCATIONS, key="inc_loc")
        with col2:
            # Hidden fields to pass the pre-selected values
            st.markdown(f"**Behaviour Type:** {behaviour_select if behaviour_select else 'Not selected'}")
            st.markdown(f"**Antecedent/Trigger:** {antecedent_select if antecedent_select else 'Not selected'}")
            # MULTIPLE INTERVENTIONS
            interventions = st.multiselect("Interventions Used *", INTERVENTIONS, key="inc_ints")
        
        duration = st.number_input("Duration (minutes) *", min_value=1, value=1, key="inc_dur")
        severity = st.slider("Severity Level (from start to end of incident) *", 1, 5, 1, key="inc_sev")
        description = st.text_area("Brief Description (Optional)", placeholder="Factual, objective description...", key="inc_desc")
        manual_critical = st.checkbox("This incident requires a Critical Incident ABCH Form (regardless of severity)", key="manual_crit")
        submitted = st.form_submit_button("Submit Incident", type="primary")
    
    if submitted:
        behaviour = behaviour_select
        antecedent = antecedent_select
        if not location or not behaviour or not antecedent or not interventions:
            st.error("Please complete all required fields marked with *")
        else:
            new_id = str(uuid.uuid4())
            is_critical = (severity >= 3) or manual_critical
            
            # Use pre-generated hypothesis or generate new one
            hyp_ai = st.session_state.get('current_hypothesis') or generate_hypothesis_ai(antecedent, behaviour, "")
            hypothesis_text = f"{hyp_ai['function']} {hyp_ai['item']}"
            
            rec = {
                "id": new_id, 
                "student_id": student_id, 
                "student_name": student["name"],
                "date": inc_date.isoformat(), 
                "time": inc_time.strftime("%H:%M:%S"),
                "day": inc_date.strftime("%A"), 
                "session": get_session_from_time(inc_time),
                "location": location, 
                "behaviour_type": behaviour, 
                "antecedent": antecedent,
                "intervention": interventions,
                "severity": severity,
                "reported_by": st.session_state.current_user["id"],
                "duration_minutes": duration, 
                "description": description or "", 
                "is_critical": is_critical,
                "hypothesis_function": hyp_ai['function'],
                "hypothesis_item": hyp_ai['item']
            }
            
            # SAVE TO DATABASE FIRST
            if save_incident_to_db(rec):
                # Then add to session state
                st.session_state.incidents.append(rec)
                st.success("‚úÖ Incident logged successfully and saved to database")
                
                if is_critical:
                    st.session_state.current_incident_id = new_id
                    st.session_state.show_critical_prompt = True
                    st.session_state.last_incident_info = {"severity": severity, "manual": manual_critical}
                    st.rerun()
            else:
                st.error("‚ùå Failed to save incident to database. Please try again.")


def render_critical_incident_page():
    """Critical Incident Form with Police Reference and Improved Labels"""
    inc_id = st.session_state.get("current_incident_id")
    quick_inc = next((i for i in st.session_state.incidents if i["id"] == inc_id), None)
    
    if not quick_inc:
        st.error("No incident found")
        return
    
    student = get_student(quick_inc["student_id"])
    st.markdown(f"## üö® Critical Incident ABCH Form")
    
    # Navigation
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("‚¨Ö Back to Students", key="back_crit_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("üè† Program Landing", key="home_crit"):
            go_to("landing")
    
    st.markdown("### Incident Details (from Quick Log)")
    with st.container(border=True):
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown(f"**Student:** {student['name']}")
            st.markdown(f"**Grade:** {student['grade']}")
        with col2:
            st.markdown(f"**Date:** {format_date_dmy(quick_inc['date'])}")
            st.markdown(f"**Time:** {format_time_12hr(quick_inc['time'])}")
        with col3:
            st.markdown(f"**Location:** {quick_inc['location']}")
            st.markdown(f"**Session:** {quick_inc['session']}")
        with col4:
            st.markdown(f"**Severity:** {quick_inc['severity']}")
        hypothesis_display = f"{quick_inc.get('hypothesis_function', '')} {quick_inc.get('hypothesis_item', '')}".strip() or 'N/A'
        st.markdown(f"**Hypothesis:** {hypothesis_display}")
    
    st.markdown("---")
    st.markdown("### ABCH Chronology")
    st.caption("Document the sequence of events. Add continuation entries if incident was prolonged.")
    
    if "abch_rows" not in st.session_state:
        st.session_state.abch_rows = []
    
    # PRIMARY ROW
    st.markdown("#### Initial Incident")
    col_header = st.columns([2, 2, 2, 2, 2])
    with col_header[0]: st.markdown("**ANTECEDENT (Triggers)**")
    with col_header[2]: st.markdown("**BEHAVIOUR**")
    with col_header[4]: st.markdown("**CONSEQUENCES**")
    
    col_subheader = st.columns([1, 1, 1, 1, 2, 2])
    with col_subheader[0]: st.caption("Location")
    with col_subheader[1]: st.caption("Context (what was happening?)")
    with col_subheader[2]: st.caption("Time")
    with col_subheader[3]: st.caption("Observed Behaviour")
    with col_subheader[4]: st.caption("What happened after?")
    with col_subheader[5]: st.caption("HYPOTHESIS (Function)")
    
    col_inputs1 = st.columns([1, 1, 1, 1, 2, 2])
    
    with col_inputs1[0]:
        location_1 = st.text_input("", value=quick_inc['location'], key="loc_1", label_visibility="collapsed")
    with col_inputs1[1]:
        context_1 = st.text_area("", placeholder="What was going on before?", 
                                key="context_1", height=100, label_visibility="collapsed")
    with col_inputs1[2]:
        time_1 = st.text_input("", value=format_time_12hr(quick_inc['time']), key="time_1", label_visibility="collapsed")
    with col_inputs1[3]:
        behaviour_1 = st.text_area("", placeholder="What did student do?", 
                                  key="behaviour_1", height=100, label_visibility="collapsed")
    with col_inputs1[4]:
        consequence_1 = st.text_area("", placeholder="Staff response? Student reaction?", 
                                    key="consequence_1", height=100, label_visibility="collapsed")
    with col_inputs1[5]:
        if "hyp_1_generated" not in st.session_state:
            st.session_state.hyp_1_generated = False
        if not st.session_state.hyp_1_generated and context_1 and behaviour_1:
            auto_hyp = generate_hypothesis(context_1, behaviour_1, consequence_1)
            st.session_state.hyp_1_auto = auto_hyp
            st.session_state.hyp_1_generated = True
        hypothesis_1 = st.text_area("", 
                                    value=st.session_state.get("hyp_1_auto", ""),
                                    placeholder="Auto-generated (editable)", 
                                    key="hypothesis_1", height=100, label_visibility="collapsed")
    
    st.markdown("---")
    
    if st.button("‚ûï Add Continuation Entry", key="add_abch_row"):
        st.session_state.abch_rows.append({})
        st.rerun()
    
    # ADDITIONAL ROWS - Changed labels
    for idx, row in enumerate(st.session_state.abch_rows):
        st.markdown(f"#### Continuation Entry {idx + 1}")
        col_add = st.columns([1, 1, 1, 1, 2, 2])
        with col_add[0]:
            row["location"] = st.text_input("", key=f"loc_{idx+2}", label_visibility="collapsed")
        with col_add[1]:
            row["context"] = st.text_area("", key=f"context_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[2]:
            row["time"] = st.text_input("", key=f"time_{idx+2}", label_visibility="collapsed")
        with col_add[3]:
            row["behaviour"] = st.text_area("", key=f"behaviour_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[4]:
            row["consequence"] = st.text_area("", key=f"consequence_{idx+2}", height=100, label_visibility="collapsed")
        with col_add[5]:
            if row.get("context") and row.get("behaviour"):
                auto_hyp_add = generate_hypothesis(row["context"], row["behaviour"], row.get("consequence", ""))
                row["hypothesis"] = st.text_area("", value=auto_hyp_add, key=f"hypothesis_{idx+2}", height=100, label_visibility="collapsed")
            else:
                row["hypothesis"] = st.text_area("", key=f"hypothesis_{idx+2}", height=100, label_visibility="collapsed")
        st.markdown("---")
    
    # INTENDED OUTCOMES with SAPOL reference field
    st.markdown("### Intended Outcomes")
    outcomes_options = [
        "Send Home", "Parent/Caregiver notified via Phone Call",
        "Student Leaving supervised areas/leaving school grounds",
        "Sexualised behaviour", "Incident ‚Äì student to student",
        "Complaint by co-located school/member of public",
        "Property damage", "Stealing", "Toileting issue",
        "ED155: Staff Injury", "ED155: Student injury",
        "Emergency services - SAPOL",
        "Emergency services - SA Ambulance",
        "Incident Internally Managed - Restorative Session",
        "Incident Internally Managed - Community Service",
        "Incident Internally Managed - Re-Entry",
        "Incident Internally Managed - Case Review",
        "Incident Internally Managed - Make-up Time"
    ]
    
    selected_outcomes = st.multiselect("Select all intended outcomes:", outcomes_options, key="intended_outcomes")
    
    # SAPOL Reference Field - triggered if SAPOL selected
    sapol_reference = ""
    if "Emergency services - SAPOL" in selected_outcomes:
        st.warning("‚ö†Ô∏è SAPOL involvement detected - Police Reference Number required")
        sapol_reference = st.text_input("SAPOL Police Reference Number *", 
                                       placeholder="Enter police reference number",
                                       key="sapol_ref")
    
    tac_notes = st.text_area("Additional Outcome Notes (e.g., TAC meeting):", 
                            placeholder="A TAC meeting will be held...",
                            key="tac_notes", height=100)
    
    st.markdown("---")
    st.markdown("### Notifications & Administration")
    col_notif1, col_notif2 = st.columns(2)
    with col_notif1:
        notified_line_manager = st.checkbox("Notified Line Manager", key="notif_manager", value=True)
        notified_parent = st.checkbox("Notified Parent/Caregiver", key="notif_parent")
    with col_notif2:
        copy_in_file = st.checkbox("Copy in student file", key="copy_file", value=True)
        safety_plan_review = st.checkbox("Safety plan review required", key="safety_review")
    
    st.markdown("---")
    st.markdown("### Staff Agreement")
    staff_name = st.session_state.current_user.get("name", "Staff Member")
    st.markdown(f"**Completing Staff Member:** {staff_name}")
    staff_agrees = st.checkbox(f"‚úì I, {staff_name}, confirm this information is accurate and complete.", 
                               key="staff_agrees")
    
    st.markdown("---")
    st.markdown("### Email Distribution")
    col_email1, col_email2 = st.columns(2)
    with col_email1:
        leader_email = st.text_input("Line Manager Email *", 
                                     value="manager@clc.sa.edu.au",
                                     key="leader_email")
    with col_email2:
        admin_email = st.text_input("Admin Email *", 
                                    value="admin@clc.sa.edu.au",
                                    key="admin_email")
    
    st.markdown("---")
    
    if st.button("üìß Submit Critical Incident Form", type="primary", use_container_width=True, key="save_crit"):
        # Validation
        errors = []
        if not context_1 or not behaviour_1 or not consequence_1 or not hypothesis_1:
            errors.append("Please complete all ABCH fields for initial incident")
        if not staff_agrees:
            errors.append("Please confirm your agreement")
        if not leader_email or "@" not in leader_email or not admin_email or "@" not in admin_email:
            errors.append("Please enter valid email addresses")
        if "Emergency services - SAPOL" in selected_outcomes and not sapol_reference:
            errors.append("SAPOL Reference Number is required when SAPOL is involved")
        
        if errors:
            for error in errors:
                st.error(f"‚ùå {error}")
        else:
            record = {
                "id": str(uuid.uuid4()),
                "created_at": datetime.now().isoformat(),
                "quick_incident_id": inc_id,
                "student_id": quick_inc["student_id"],
                "student_name": student["name"],
                "incident_type": "Critical",
                "ABCH_primary": {
                    "location": location_1,
                    "context": context_1,
                    "time": time_1,
                    "behaviour": behaviour_1,
                    "consequence": consequence_1,
                    "hypothesis": hypothesis_1
                },
                "ABCH_additional": st.session_state.abch_rows.copy(),
                "intended_outcomes": selected_outcomes,
                "sapol_reference": sapol_reference if sapol_reference else None,
                "tac_notes": tac_notes,
                "notifications": {
                    "line_manager": notified_line_manager,
                    "parent": notified_parent,
                    "copy_in_file": copy_in_file,
                    "safety_plan_review": safety_plan_review
                },
                "staff_agreement": {
                    "staff_name": staff_name,
                    "agreed": staff_agrees,
                    "timestamp": datetime.now().isoformat()
                },
                "leader_email": leader_email,
                "admin_email": admin_email
            }
            
            st.session_state.critical_incidents.append(record)
            st.session_state.abch_rows = []
            st.session_state.hyp_1_generated = False
            
            st.success("‚úÖ Critical incident form saved to database")
            
            # GENERATE ADMIN SUMMARY
            admin_summary = generate_admin_summary(record, student, staff_name)
            
            # Show admin summary
            with st.expander("üìã ADMIN SUMMARY (For External Incident Log)", expanded=True):
                st.text_area("Copy this summary for departmental log:", admin_summary, height=400, key="admin_summary_display")
                st.download_button(
                    "üì• Download Admin Summary",
                    admin_summary,
                    file_name=f"Admin_Summary_{student['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt",
                    mime="text/plain"
                )
            
            # SEND EMAILS
            staff_email = st.session_state.current_user.get("email", "staff@example.com")
            send_critical_incident_email(record, student, staff_email, leader_email, admin_email)
            
            st.markdown("---")
            st.info("‚úâÔ∏è Emails sent to Line Manager, Admin, and completing staff member")
            st.info("üíæ Critical incident data saved in student's file")
            st.info("üìã Admin summary generated for external log")
            if sapol_reference:
                st.info(f"üöî SAPOL Reference: {sapol_reference}")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("üìä View Analysis", type="primary", use_container_width=True, key="view_analysis"):
                    go_to("student_analysis", selected_student_id=quick_inc["student_id"])
            with col2:
                if st.button("‚Ü©Ô∏è Back to Students", use_container_width=True, key="back_crit_after"):
                    go_to("program_students", selected_program=student["program"])
            with col3:
                if st.button("üè† Program Landing", use_container_width=True, key="home_crit_after"):
                    go_to("landing")
def render_student_analysis_page():
    """Comprehensive Data Analysis with Berry Street Education Model"""
    student_id = st.session_state.get("selected_student_id")
    student = get_student(student_id)
    if not student:
        st.error("No student selected")
        return
    
    st.markdown(f"## üìä Comprehensive Behaviour Analysis ‚Äî {student['name']}")
    st.caption("Evidence-based analysis prepared by Learning and Behaviour Unit")
    st.caption("Using ABA, Trauma-Informed Practice, Berry Street Education Model, and CPI principles")
    
    # Display placement information
    if student.get('placement_start'):
        try:
            start_dt = datetime.fromisoformat(student['placement_start'])
            if student.get('placement_end'):
                end_dt = datetime.fromisoformat(student['placement_end'])
                status_txt = "Completed"
            else:
                end_dt = datetime.now()
                status_txt = "Ongoing"
            
            days_enrolled = (end_dt.date() - start_dt.date()).days
            
            st.markdown("---")
            pcol1, pcol2, pcol3 = st.columns(3)
            with pcol1:
                st.metric("Placement Start", start_dt.strftime('%d/%m/%Y'))
            with pcol2:
                if student.get('placement_end'):
                    st.metric("Placement End", datetime.fromisoformat(student['placement_end']).strftime('%d/%m/%Y'))
                else:
                    st.metric("Status", status_txt)
            with pcol3:
                st.metric("Days Enrolled", days_enrolled)
            st.markdown("---")
        except:
            pass
    
    # Navigation
    col1, col2 = st.columns([3, 1])
    with col1:
        if st.button("‚¨Ö Back to Students", key="back_analysis_top"):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("üè† Program Landing", key="home_analysis"):
            go_to("landing")
    
    quick = [i for i in st.session_state.incidents if i["student_id"] == student_id]
    crit = [c for c in st.session_state.critical_incidents if c["student_id"] == student_id]
    
    if not quick and not crit:
        st.info("No incident data available yet.")
        if st.button("‚Ü©Ô∏è Back", key="back_no_data"):
            go_to("program_students", selected_program=student["program"])
        return
    
    # Build dataframe
    quick_df = pd.DataFrame(quick) if quick else pd.DataFrame()
    crit_df = pd.DataFrame(crit) if crit else pd.DataFrame()
    
    if not quick_df.empty:
        quick_df["incident_type"] = "Quick"
        quick_df["date_parsed"] = pd.to_datetime(quick_df["date"])
        if "intervention" in quick_df.columns:
            quick_df["intervention_str"] = quick_df["intervention"].apply(lambda x: ", ".join(x) if isinstance(x, list) else str(x))
    
    if not crit_df.empty:
        crit_df["incident_type"] = "Critical"
        crit_df["date_parsed"] = pd.to_datetime(crit_df.get("created_at", datetime.now().isoformat()))
        crit_df["severity"] = 5
        crit_df["antecedent"] = crit_df["ABCH_primary"].apply(lambda d: d.get("context","") if isinstance(d, dict) else "")
        crit_df["behaviour_type"] = crit_df["ABCH_primary"].apply(lambda d: d.get("behaviour","") if isinstance(d, dict) else "")
    
    full_df = pd.concat([quick_df, crit_df], ignore_index=True).sort_values("date_parsed")
    full_df["hour"] = pd.to_datetime(full_df["time"], format="%H:%M:%S", errors="coerce").dt.hour
    full_df["day_of_week"] = full_df["date_parsed"].dt.day_name()
    
    # Prepare split dataframes
    quick_only_df = full_df[full_df['incident_type'] == 'Quick'].copy() if 'Quick' in full_df['incident_type'].values else pd.DataFrame()
    crit_only_df = full_df[full_df['incident_type'] == 'Critical'].copy() if 'Critical' in full_df['incident_type'].values else pd.DataFrame()
    
    # OVERVIEW
    st.markdown("### üìà Executive Summary")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1: st.metric("Total", len(full_df))
    with col2: st.metric("Critical", len(full_df[full_df["incident_type"] == "Critical"]))
    with col3: st.metric("Avg Severity", f"{full_df['severity'].mean():.1f}")
    with col4:
        days = max((full_df["date_parsed"].max() - full_df["date_parsed"].min()).days, 1)
        st.metric("Days Span", days)
    with col5:
        st.metric("Per Day", f"{len(full_df) / days:.1f}")
    
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 1: DAILY FREQUENCY
    # ================================================================
    
    st.markdown("### üìÖ Daily Incident Frequency - Regular vs Critical")
    st.caption("Understanding the relationship between regular incidents and critical escalations")
    
    fig1 = go.Figure()
    
    if not quick_only_df.empty:
        daily_quick = quick_only_df.groupby(quick_only_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig1.add_trace(go.Bar(
            x=daily_quick["date_parsed"], 
            y=daily_quick["count"],
            name='Regular Incidents',
            marker=dict(color='#4A90A4', line=dict(color='white', width=1)),
            text=daily_quick["count"],
            textposition='inside',
            textfont=dict(color='white', size=11, family='Arial Black'),
            hovertemplate='<b>Date:</b> %{x}<br><b>Regular:</b> %{y}<extra></extra>'
        ))
    
    if not crit_only_df.empty:
        daily_crit = crit_only_df.groupby(crit_only_df["date_parsed"].dt.date).size().reset_index(name="count")
        fig1.add_trace(go.Bar(
            x=daily_crit["date_parsed"], 
            y=daily_crit["count"],
            name='Critical Incidents',
            marker=dict(color='#ef4444', line=dict(color='white', width=1)),
            text=daily_crit["count"],
            textposition='inside',
            textfont=dict(color='white', size=11, family='Arial Black'),
            hovertemplate='<b>Date:</b> %{x}<br><b>Critical:</b> %{y}<extra></extra>'
        ))
    
    fig1.update_layout(
        height=350,
        barmode='stack',
        xaxis_title="<b>Date</b>",
        yaxis_title="<b>Incident Count</b>",
        plot_bgcolor='#f8fafc',
        paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        yaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', gridwidth=1, showline=True, linewidth=2, linecolor='#cbd5e1'),
        xaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1', gridcolor='#e2e8f0'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1),
        hovermode='x unified'
    )
    st.plotly_chart(fig1, use_container_width=True)
    
    if not quick_only_df.empty and not crit_only_df.empty:
        escalation_rate = (len(crit_only_df) / len(quick_only_df)) * 100 if len(quick_only_df) > 0 else 0
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Regular Incidents", len(quick_only_df), help="Severity 1-2")
        with col2:
            st.metric("Critical Incidents", len(crit_only_df), help="Severity 3+")
        with col3:
            st.metric("Escalation Rate", f"{escalation_rate:.1f}%", help="% of incidents that escalate to critical")
    
    with st.expander("üí° Clinical Interpretation (Berry Street Body Domain)"):
        st.markdown(
            "**Pattern Recognition:** Days with multiple regular incidents may predict critical escalation. "
            "When you see clustering of regular incidents, it indicates the student's nervous system is already dysregulated.\n\n"
            "**Berry Street Body:** On high-frequency days, increase proactive regulation - breathing exercises, "
            "movement breaks, sensory activities. The goal is to widen the Window of Tolerance before it narrows further.\n\n"
            "**Prevention Strategy:** If you see 2+ regular incidents in one day, immediately implement intensive Body domain supports "
            "to prevent critical escalation."
        )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 2: BEHAVIOURS
    # ================================================================
    
    st.markdown("### üéØ Behaviour Types - Regular vs Critical")
    st.caption("Which behaviours escalate to critical incidents?")
    
    all_behaviours = full_df["behaviour_type"].value_counts().head(6).index.tolist()
    
    quick_beh_counts = []
    crit_beh_counts = []
    for beh in all_behaviours:
        quick_count = len(quick_only_df[quick_only_df["behaviour_type"] == beh]) if not quick_only_df.empty else 0
        crit_count = len(crit_only_df[crit_only_df["behaviour_type"] == beh]) if not crit_only_df.empty else 0
        quick_beh_counts.append(quick_count)
        crit_beh_counts.append(crit_count)
    
    fig2 = go.Figure()
    
    fig2.add_trace(go.Bar(
        y=all_behaviours, x=quick_beh_counts, name='Regular', orientation='h',
        marker=dict(color='#4A90A4', line=dict(color='white', width=1)),
        text=quick_beh_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black'),
        hovertemplate='<b>%{y}</b><br>Regular: %{x}<extra></extra>'
    ))
    
    fig2.add_trace(go.Bar(
        y=all_behaviours, x=crit_beh_counts, name='Critical', orientation='h',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_beh_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black'),
        hovertemplate='<b>%{y}</b><br>Critical: %{x}<extra></extra>'
    ))
    
    fig2.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Incident Count</b>", yaxis_title="",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        xaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', gridwidth=1, showline=True, linewidth=2, linecolor='#cbd5e1'),
        yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1),
        hovermode='y unified'
    )
    st.plotly_chart(fig2, use_container_width=True)
    
    st.markdown("**Escalation Risk by Behaviour:**")
    risk_data = []
    for i, beh in enumerate(all_behaviours):
        total = quick_beh_counts[i] + crit_beh_counts[i]
        if total > 0:
            risk_pct = (crit_beh_counts[i] / total) * 100
            risk_level = "üî¥ High Risk" if risk_pct > 50 else "üü° Medium Risk" if risk_pct > 25 else "üü¢ Low Risk"
            risk_data.append(f"‚Ä¢ **{beh}**: {risk_pct:.0f}% escalate to critical - {risk_level}")
    
    for risk_item in risk_data:
        st.markdown(risk_item)
    
    with st.expander("üí° Clinical Interpretation (Behaviour as Communication)"):
        if all_behaviours:
            primary_beh = all_behaviours[0]
            st.markdown(
                f"**Primary Concern:** {primary_beh} is the most common behaviour. "
                "Behaviours with high escalation rates (>50%) need immediate intervention planning.\n\n"
                "**Behaviour Analysis:** Focus on high-risk behaviours first. If a behaviour frequently escalates, "
                "it means the student lacks skills or supports to regulate at that level.\n\n"
                "**Berry Street:** All behaviour is communication. High escalation rates tell us the student needs:\n"
                "1. **BODY**: More regulation tools before the behaviour occurs\n"
                "2. **RELATIONSHIP**: Stronger connection to trusted adults who can co-regulate\n"
                "3. **STAMINA**: Skill-building for persistence through challenges"
            )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 3: TRIGGERS
    # ================================================================
    
    st.markdown("### üîç Trigger Analysis - What Leads to Critical Escalation?")
    st.caption("Understanding which antecedents most often result in critical incidents")
    
    all_triggers = full_df["antecedent"].value_counts().head(6).index.tolist()
    
    quick_ant_counts = []
    crit_ant_counts = []
    for ant in all_triggers:
        quick_count = len(quick_only_df[quick_only_df["antecedent"] == ant]) if not quick_only_df.empty else 0
        crit_count = len(crit_only_df[crit_only_df["antecedent"] == ant]) if not crit_only_df.empty else 0
        quick_ant_counts.append(quick_count)
        crit_ant_counts.append(crit_count)
    
    fig3 = go.Figure()
    
    fig3.add_trace(go.Bar(
        y=all_triggers, x=quick_ant_counts, name='Regular', orientation='h',
        marker=dict(color='#4A90A4', line=dict(color='white', width=1)),
        text=quick_ant_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig3.add_trace(go.Bar(
        y=all_triggers, x=crit_ant_counts, name='Critical', orientation='h',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_ant_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig3.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Incident Count</b>", yaxis_title="",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        xaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1)
    )
    st.plotly_chart(fig3, use_container_width=True)
    
    st.markdown("**Critical Escalation Risk by Trigger:**")
    for i, ant in enumerate(all_triggers):
        total = quick_ant_counts[i] + crit_ant_counts[i]
        if total > 0:
            risk_pct = (crit_ant_counts[i] / total) * 100
            color = "üî¥" if risk_pct > 60 else "üü°" if risk_pct > 30 else "üü¢"
            st.markdown(f"{color} **{ant}**: {risk_pct:.0f}% lead to critical incidents ({crit_ant_counts[i]}/{total})")
    
    with st.expander("üí° Clinical Interpretation (Proactive Prevention)"):
        st.markdown(
            "**High-Risk Triggers (>60%):** These antecedents almost always lead to crisis. Implement intensive preventative supports:\n"
            "‚Ä¢ Modify environment to prevent trigger\n"
            "‚Ä¢ Pre-teach coping strategies specific to this trigger\n"
            "‚Ä¢ Provide extra staff support when trigger is likely\n\n"
            "**Medium-Risk Triggers (30-60%):** Student has some capacity but needs support:\n"
            "‚Ä¢ Berry Street STAMINA: Build persistence through small exposures\n"
            "‚Ä¢ Provide choice and control\n"
            "‚Ä¢ Co-regulation available immediately\n\n"
            "**Low-Risk Triggers (<30%):** Student managing well with current supports - maintain strategies."
        )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 4: SEVERITY WITH ZONES
    # ================================================================
    
    st.markdown("### üìä Severity Trajectory - Window of Tolerance Analysis")
    st.caption("Tracking whether incidents are getting more or less severe over time")
    
    fig4 = go.Figure()
    
    fig4.add_hrect(y0=0, y1=2.5, fillcolor="#d1fae5", opacity=0.2, 
                   annotation_text="Within Window (Regulation)", annotation_position="top left", line_width=0)
    fig4.add_hrect(y0=2.5, y1=3.5, fillcolor="#fef3c7", opacity=0.2,
                   annotation_text="Edge of Window", annotation_position="top left", line_width=0)
    fig4.add_hrect(y0=3.5, y1=5.5, fillcolor="#fee2e2", opacity=0.2,
                   annotation_text="Outside Window (Crisis)", annotation_position="top left", line_width=0)
    
    if not quick_only_df.empty:
        fig4.add_trace(go.Scatter(
            x=quick_only_df["date_parsed"], y=quick_only_df["severity"],
            mode='markers', name='Regular',
            marker=dict(size=12, color='#4A90A4', opacity=0.7, line=dict(color='white', width=1.5)),
            hovertemplate='<b>Date:</b> %{x}<br><b>Severity:</b> %{y}<extra></extra>'
        ))
    
    if not crit_only_df.empty:
        fig4.add_trace(go.Scatter(
            x=crit_only_df["date_parsed"], y=crit_only_df["severity"],
            mode='markers', name='Critical',
            marker=dict(size=15, color='#ef4444', opacity=0.8, symbol='diamond', line=dict(color='white', width=2)),
            hovertemplate='<b>Date:</b> %{x}<br><b>Severity:</b> %{y}<extra></extra>'
        ))
    
    if len(full_df) >= 3:
        z = np.polyfit(range(len(full_df)), full_df["severity"], 1)
        p = np.poly1d(z)
        trend_color = '#22c55e' if z[0] < 0 else '#ef4444'
        fig4.add_trace(go.Scatter(
            x=full_df["date_parsed"], y=p(range(len(full_df))),
            mode='lines', name='Trend',
            line=dict(color=trend_color, width=3, dash='dash'),
            hovertemplate='Trend<extra></extra>'
        ))
    
    fig4.update_layout(
        height=400, xaxis_title="<b>Date</b>", yaxis_title="<b>Severity Level</b>",
        yaxis=dict(range=[0, 5.5], tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        xaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1', gridcolor='#e2e8f0'),
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.9)', bordercolor='#cbd5e1', borderwidth=1),
        hovermode='closest'
    )
    st.plotly_chart(fig4, use_container_width=True)
    
    if len(full_df) >= 5:
        recent_avg = full_df.tail(5)["severity"].mean()
        early_avg = full_df.head(5)["severity"].mean()
        trend_dir = "improving" if recent_avg < early_avg else "worsening" if recent_avg > early_avg else "stable"
        trend_emoji = "üìà" if trend_dir == "improving" else "üìâ" if trend_dir == "worsening" else "‚û°Ô∏è"
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Early Average", f"{early_avg:.1f}", help="First 5 incidents")
        with col2:
            st.metric("Recent Average", f"{recent_avg:.1f}", help="Last 5 incidents")
        with col3:
            st.metric("Trend", f"{trend_emoji} {trend_dir.title()}")
    
    with st.expander("üí° Clinical Interpretation (Window of Tolerance)"):
        st.markdown(
            "**Window of Tolerance (Siegel, 1999):** The optimal arousal zone where the student can think, learn, and regulate.\n\n"
            "**Green Zone (Severity 1-2):** Student is within or near their window. Accessible to support.\n\n"
            "**Yellow Zone (Severity 2.5-3.5):** Student is at the edge. CO-REGULATION NEEDED NOW.\n\n"
            "**Red Zone (Severity 3.5-5):** Outside window in survival mode. Safety first, teach later."
        )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 5: LOCATION
    # ================================================================
    
    st.markdown("### üìç Location Hotspots - Where Do Critical Incidents Occur?")
    st.caption("Environmental factors and escalation patterns by location")
    
    all_locations = full_df["location"].value_counts().head(6).index.tolist()
    
    quick_loc_counts = []
    crit_loc_counts = []
    for loc in all_locations:
        quick_count = len(quick_only_df[quick_only_df["location"] == loc]) if not quick_only_df.empty else 0
        crit_count = len(crit_only_df[crit_only_df["location"] == loc]) if not crit_only_df.empty else 0
        quick_loc_counts.append(quick_count)
        crit_loc_counts.append(crit_count)
    
    fig5 = go.Figure()
    
    fig5.add_trace(go.Bar(
        y=all_locations, x=quick_loc_counts, name='Regular', orientation='h',
        marker=dict(color='#4A90A4', line=dict(color='white', width=1)),
        text=quick_loc_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig5.add_trace(go.Bar(
        y=all_locations, x=crit_loc_counts, name='Critical', orientation='h',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_loc_counts, textposition='inside',
        textfont=dict(color='white', size=11, family='Arial Black')
    ))
    
    fig5.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Incident Count</b>", yaxis_title="",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        xaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1)
    )
    st.plotly_chart(fig5, use_container_width=True)
    
    st.markdown("**Environmental Risk Assessment:**")
    for i, loc in enumerate(all_locations):
        total = quick_loc_counts[i] + crit_loc_counts[i]
        if total > 0:
            risk_pct = (crit_loc_counts[i] / total) * 100
            if risk_pct > 60:
                risk_level = "üî¥ HIGH RISK ENVIRONMENT"
                recommendation = "Immediate environmental modification needed"
            elif risk_pct > 30:
                risk_level = "üü° MODERATE RISK"
                recommendation = "Enhanced supervision and supports"
            else:
                risk_level = "üü¢ MANAGED ENVIRONMENT"
                recommendation = "Current strategies effective"
            st.markdown(f"**{loc}**: {risk_pct:.0f}% escalate - {risk_level} - *{recommendation}*")
    
    with st.expander("üí° Clinical Interpretation (Environmental Strategies)"):
        if all_locations:
            st.markdown(
                f"**Primary Hotspot:** {all_locations[0]}\n\n"
                "**Berry Street BODY - Sensory Environment:**\n"
                "‚Ä¢ Lighting: Consider natural light or lamps\n"
                "‚Ä¢ Noise: Provide noise-cancelling headphones or quiet spaces\n"
                "‚Ä¢ Space: Create clear pathways and defined areas"
            )
    st.markdown("---")
    
    # ================================================================
    # ENHANCED GRAPH 6: TIME OF DAY
    # ================================================================
    
    st.markdown("### ‚è∞ Time of Day Patterns - When Does Escalation Occur?")
    st.caption("Understanding daily rhythm and predicting high-risk periods")
    
    session_order = ['Morning', 'Middle', 'Afternoon']
    
    quick_session_counts = []
    crit_session_counts = []
    for session in session_order:
        quick_count = len(quick_only_df[quick_only_df["session"] == session]) if not quick_only_df.empty and 'session' in quick_only_df.columns else 0
        crit_count = len(crit_only_df[crit_only_df["session"] == session]) if not crit_only_df.empty and 'session' in crit_only_df.columns else 0
        quick_session_counts.append(quick_count)
        crit_session_counts.append(crit_count)
    
    fig6 = go.Figure()
    
    fig6.add_trace(go.Bar(
        x=session_order, y=quick_session_counts, name='Regular',
        marker=dict(color='#4A90A4', line=dict(color='white', width=1)),
        text=quick_session_counts, textposition='inside',
        textfont=dict(color='white', size=12, family='Arial Black')
    ))
    
    fig6.add_trace(go.Bar(
        x=session_order, y=crit_session_counts, name='Critical',
        marker=dict(color='#ef4444', line=dict(color='white', width=1)),
        text=crit_session_counts, textposition='inside',
        textfont=dict(color='white', size=12, family='Arial Black')
    ))
    
    fig6.update_layout(
        height=350, barmode='stack', xaxis_title="<b>Time of Day</b>", yaxis_title="<b>Incident Count</b>",
        plot_bgcolor='#f8fafc', paper_bgcolor='white',
        font=dict(color='#0f172a', size=12, family='Arial'),
        yaxis=dict(tickmode='linear', tick0=0, dtick=1, gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
        xaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1'),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor='rgba(255,255,255,0.8)', bordercolor='#cbd5e1', borderwidth=1)
    )
    st.plotly_chart(fig6, use_container_width=True)
    
    st.markdown("**Escalation Risk by Time of Day:**")
    for i, session in enumerate(session_order):
        total = quick_session_counts[i] + crit_session_counts[i]
        if total > 0:
            risk_pct = (crit_session_counts[i] / total) * 100
            risk_emoji = "üî¥" if risk_pct > 50 else "üü°" if risk_pct > 25 else "üü¢"
            st.markdown(f"{risk_emoji} **{session}**: {risk_pct:.0f}% escalate to critical ({total} total incidents)")
    
    with st.expander("üí° Clinical Interpretation (Circadian Regulation)"):
        if quick_session_counts:
            peak_session = session_order[quick_session_counts.index(max(quick_session_counts))] if max(quick_session_counts) > 0 else "Unknown"
            st.markdown(
                f"**Peak Incident Time:** {peak_session}\n\n"
                "**Berry Street BODY:** Proactive regulation before peak periods - breathing, movement, sensory check-ins."
            )
    st.markdown("---")
    
    # QUICK VS CRITICAL COMPARISON SECTION
    st.markdown("## üìä Quick vs Critical Incident Analysis")
    st.caption("Understanding the relationship between quick logs and critical incidents helps identify escalation patterns")
    
    quick_only = [i for i in quick if not i.get("is_critical")]
    critical_data = crit
    
    col_q, col_c = st.columns(2)
    with col_q:
        st.metric("Quick Incidents", len(quick_only), help="Standard behaviour logs (Severity 1-2)")
    with col_c:
        st.metric("Critical Incidents", len(critical_data), help="Severity 3+ requiring ABCH form")
    
    st.markdown("---")
    
    # Day of Week
    st.markdown("### üìÜ Day of Week Patterns")
    day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    day_counts = full_df["day_of_week"].value_counts().reindex(day_order, fill_value=0)
    fig7 = go.Figure()
    fig7.add_trace(go.Bar(
        x=day_counts.index, y=day_counts.values,
        marker=dict(color='#008080'),
        text=day_counts.values, textposition='outside'
    ))
    fig7.update_layout(
        height=280, showlegend=False, yaxis_title="Total",
        plot_bgcolor='white', paper_bgcolor='white',
        font=dict(color='#334155', size=11)
    )
    st.plotly_chart(fig7, use_container_width=True)
    
    high_day = day_counts.idxmax()
    with st.expander("üí° Clinical Interpretation (Berry Street Relationship)"):
        st.markdown(f"**{high_day}** has most incidents. Consider connection routines. " +
                   "**Berry Street Relationship:** Strong connections reduce incidents.")
    st.markdown("---")
    
    # ================================================================
    # NEW: HYPOTHESIS/FUNCTION ANALYSIS
    # ================================================================
    
    st.markdown("### üéØ Behaviour Function Analysis")
    st.caption("Understanding the purpose behind behaviours using Applied Behaviour Analysis (ABA)")
    
    # Analyse hypothesis patterns
    if not quick_only_df.empty and 'hypothesis_function' in quick_only_df.columns:
        # Function distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Function Distribution**")
            func_counts = quick_only_df['hypothesis_function'].value_counts()
            
            fig_func = go.Figure()
            colors_func = ['#008080' if f == 'To avoid' else '#20B2AA' for f in func_counts.index]
            fig_func.add_trace(go.Pie(
                labels=func_counts.index,
                values=func_counts.values,
                marker=dict(colors=colors_func, line=dict(color='white', width=2)),
                textinfo='label+percent',
                textfont=dict(size=12, family='Arial'),
                hole=0.4
            ))
            fig_func.update_layout(
                height=280,
                showlegend=False,
                margin=dict(t=20, b=20, l=20, r=20),
                paper_bgcolor='white'
            )
            st.plotly_chart(fig_func, use_container_width=True)
        
        with col2:
            st.markdown("**What Student is Seeking**")
            if 'hypothesis_item' in quick_only_df.columns:
                item_counts = quick_only_df['hypothesis_item'].value_counts()
                
                fig_item = go.Figure()
                fig_item.add_trace(go.Bar(
                    x=item_counts.values,
                    y=item_counts.index,
                    orientation='h',
                    marker=dict(color=['#4A90A4', '#6BB9A0', '#48D1CC', '#7FFFD4'][:len(item_counts)],
                               line=dict(color='white', width=1)),
                    text=item_counts.values,
                    textposition='inside',
                    textfont=dict(color='white', size=11, family='Arial Black')
                ))
                fig_item.update_layout(
                    height=280,
                    xaxis_title="<b>Count</b>",
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    margin=dict(t=20, b=40, l=20, r=20),
                    xaxis=dict(gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
                    yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1')
                )
                st.plotly_chart(fig_item, use_container_width=True)
        
        # Interpretation
        primary_function = func_counts.index[0] if len(func_counts) > 0 else "Unknown"
        primary_item = item_counts.index[0] if 'hypothesis_item' in quick_only_df.columns and len(item_counts) > 0 else "Unknown"
        
        with st.expander("üí° Clinical Interpretation (Behaviour Function)"):
            st.markdown(f"""
**Primary Hypothesis:** {primary_function} {primary_item}

**What This Means:**
- **{primary_function}** behaviours indicate the student is trying to {'escape or avoid something aversive' if primary_function == 'To avoid' else 'access or obtain something desired'}
- **{primary_item}** focus suggests intervention should target {'reducing demands or modifying environment' if primary_item == 'Activity' else 'sensory regulation strategies' if primary_item == 'Sensory' else 'increasing positive attention and connection' if primary_item == 'Attention' else 'teaching appropriate requesting skills'}

**Berry Street Connection:**
- {'**BODY Domain:** Focus on regulation before challenging activities' if primary_function == 'To avoid' else '**RELATIONSHIP Domain:** Increase positive attention and connection opportunities'}
            """)
    
    st.markdown("---")
    
    # ================================================================
    # NEW: AUSTRALIAN CURRICULUM CAPABILITIES SECTION
    # ================================================================
    
    st.markdown("### üéì Australian Curriculum - General Capabilities Analysis")
    st.caption("Linking behaviour support to curriculum outcomes (Personal and Social Capability focus)")
    
    top_beh = full_df["behaviour_type"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_ant = full_df["antecedent"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_loc = full_df["location"].mode()[0] if len(full_df) > 0 else "Unknown"
    top_session = full_df["session"].mode()[0] if len(full_df) > 0 else "Unknown"
    
    # Get AC capability information
    ac_info = get_ac_capability_for_behaviour(top_beh)
    ant_ac_info = get_ac_skills_for_antecedent(top_ant)
    
    if ac_info:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.markdown(f"""
<div style='background: linear-gradient(135deg, #E8F4F8 0%, #D4EEF4 100%); padding: 1.5rem; border-radius: 10px; border-left: 4px solid {ac_info['color']}; margin-bottom: 1rem;'>
    <h4 style='color: #008080; margin: 0 0 0.5rem 0;'>Primary Capability: {ac_info['name']}</h4>
    <p style='margin: 0; color: #333;'><strong>Focus Elements:</strong> {', '.join(ac_info['elements'])}</p>
</div>
            """, unsafe_allow_html=True)
            
            st.markdown("**Skills to Develop (AC Aligned):**")
            for skill in ac_info['skills_to_develop']:
                st.markdown(f"‚Ä¢ {skill}")
            
            st.markdown("**Curriculum Descriptors:**")
            for desc in ac_info['ac_descriptors']:
                st.markdown(f"‚Ä¢ _{desc}_")
        
        with col2:
            # AC Capability visual
            st.markdown("**Capability Focus**")
            
            # Create a simple gauge/indicator
            fig_ac = go.Figure()
            
            capabilities = ["PSC", "CCT", "EU", "ICU"]
            cap_names = ["Personal & Social", "Critical Thinking", "Ethical", "Intercultural"]
            cap_colors = ['#4A90A4', '#6BB9A0', '#E8B960', '#D4A574']
            
            # Highlight the primary capability
            values = [1 if c == ac_info['code'] else 0.3 for c in capabilities]
            
            fig_ac.add_trace(go.Bar(
                x=cap_names,
                y=values,
                marker=dict(color=cap_colors, line=dict(color='white', width=1)),
                text=['PRIMARY' if v == 1 else '' for v in values],
                textposition='inside',
                textfont=dict(color='white', size=10, family='Arial Black')
            ))
            fig_ac.update_layout(
                height=200,
                showlegend=False,
                yaxis=dict(visible=False),
                xaxis=dict(tickangle=-45),
                margin=dict(t=10, b=60, l=10, r=10),
                plot_bgcolor='white',
                paper_bgcolor='white'
            )
            st.plotly_chart(fig_ac, use_container_width=True)
    
    # Generate AC-aligned learning goals
    learning_goals = generate_ac_learning_goals(top_beh, top_ant, student['grade'])
    
    st.markdown("**Suggested Learning Goals (Grade-Appropriate):**")
    goals_table = []
    for goal in learning_goals:
        goals_table.append({
            "Skill Focus": goal['skill'],
            "Learning Goal": goal['goal'],
            "AC Descriptor": goal['ac_descriptor'],
            "Complexity": goal['complexity'].title()
        })
    
    if goals_table:
        goals_df = pd.DataFrame(goals_table)
        st.dataframe(goals_df, hide_index=True, use_container_width=True)
    
    with st.expander("üí° Understanding AC General Capabilities"):
        st.markdown("""
**Personal and Social Capability** is developed when students:
- Learn to understand themselves and others
- Manage their emotions, relationships and lives
- Develop resilience and a sense of self-worth
- Work effectively with others
- Make responsible decisions

**Connection to Behaviour Support:**
All behaviour is an attempt to meet a need. When we support students to develop Personal and Social Capability skills, we give them the tools to meet their needs in prosocial ways. This is not just behaviour management - it is curriculum-aligned skill development.

**Assessment Note:** Progress in these capabilities should be documented alongside behaviour data to show growth in underlying skills, not just reduction in incidents.
        """)
    
    st.markdown("---")
    
    # ================================================================
    # NEW: INTERVENTION EFFECTIVENESS ANALYSIS
    # ================================================================
    
    st.markdown("### üíä Intervention Analysis")
    st.caption("Which strategies are being used and their connection to capability development")
    
    if not quick_only_df.empty and 'intervention' in quick_only_df.columns:
        # Flatten interventions
        all_interventions = []
        for interventions in quick_only_df['intervention']:
            if isinstance(interventions, list):
                all_interventions.extend(interventions)
            elif isinstance(interventions, str):
                all_interventions.append(interventions)
        
        if all_interventions:
            intervention_counts = pd.Series(all_interventions).value_counts()
            
            col1, col2 = st.columns([3, 2])
            
            with col1:
                fig_int = go.Figure()
                fig_int.add_trace(go.Bar(
                    y=intervention_counts.index[:8],
                    x=intervention_counts.values[:8],
                    orientation='h',
                    marker=dict(
                        color=['#008080', '#20B2AA', '#48D1CC', '#4A90A4', '#6BB9A0', '#7FFFD4', '#40E0D0', '#00CED1'][:len(intervention_counts)],
                        line=dict(color='white', width=1)
                    ),
                    text=intervention_counts.values[:8],
                    textposition='inside',
                    textfont=dict(color='white', size=11, family='Arial Black')
                ))
                fig_int.update_layout(
                    height=300,
                    xaxis_title="<b>Times Used</b>",
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    margin=dict(t=20, b=40, l=20, r=20),
                    xaxis=dict(gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
                    yaxis=dict(showline=True, linewidth=2, linecolor='#cbd5e1')
                )
                st.plotly_chart(fig_int, use_container_width=True)
            
            with col2:
                st.markdown("**Intervention-Capability Alignment**")
                
                # Get AC alignment for top interventions
                top_interventions = intervention_counts.head(4).index.tolist()
                alignments = get_intervention_ac_alignment(top_interventions)
                
                for align in alignments:
                    st.markdown(f"""
<div style='background: #f8f9fa; padding: 0.8rem; border-radius: 6px; margin-bottom: 0.5rem; border-left: 3px solid #008080;'>
    <strong style='color: #008080;'>{align['intervention']}</strong><br>
    <small style='color: #666;'>{align['capability']} - {align['element']}</small><br>
    <small style='color: #333;'>Supports: {align['supports']}</small>
</div>
                    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # ================================================================
    # NEW: WEEK-OVER-WEEK TREND ANALYSIS
    # ================================================================
    
    st.markdown("### üìà Weekly Trend Analysis")
    st.caption("Tracking progress over time to inform intervention adjustments")
    
    if len(full_df) >= 7:
        full_df['week'] = full_df['date_parsed'].dt.isocalendar().week
        weekly_stats = full_df.groupby('week').agg({
            'severity': ['count', 'mean'],
            'incident_type': lambda x: (x == 'Critical').sum()
        }).reset_index()
        weekly_stats.columns = ['Week', 'Total Incidents', 'Avg Severity', 'Critical Count']
        
        if len(weekly_stats) >= 2:
            col1, col2 = st.columns(2)
            
            with col1:
                fig_weekly = go.Figure()
                fig_weekly.add_trace(go.Scatter(
                    x=weekly_stats['Week'],
                    y=weekly_stats['Total Incidents'],
                    mode='lines+markers',
                    name='Total Incidents',
                    line=dict(color='#008080', width=3),
                    marker=dict(size=10, color='#008080', line=dict(color='white', width=2))
                ))
                fig_weekly.add_trace(go.Scatter(
                    x=weekly_stats['Week'],
                    y=weekly_stats['Critical Count'],
                    mode='lines+markers',
                    name='Critical Incidents',
                    line=dict(color='#DC3545', width=3, dash='dash'),
                    marker=dict(size=10, color='#DC3545', symbol='diamond', line=dict(color='white', width=2))
                ))
                fig_weekly.update_layout(
                    height=280,
                    xaxis_title="<b>Week Number</b>",
                    yaxis_title="<b>Count</b>",
                    plot_bgcolor='white',
                    paper_bgcolor='white',
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                    xaxis=dict(gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1'),
                    yaxis=dict(gridcolor='#e2e8f0', showline=True, linewidth=2, linecolor='#cbd5e1')
                )
                st.plotly_chart(fig_weekly, use_container_width=True)
            
            with col2:
                # Week-over-week change
                if len(weekly_stats) >= 2:
                    last_week = weekly_stats.iloc[-1]
                    prev_week = weekly_stats.iloc[-2]
                    
                    incident_change = last_week['Total Incidents'] - prev_week['Total Incidents']
                    severity_change = last_week['Avg Severity'] - prev_week['Avg Severity']
                    
                    st.markdown("**Week-over-Week Change**")
                    
                    col_a, col_b = st.columns(2)
                    with col_a:
                        delta_color = "inverse" if incident_change <= 0 else "normal"
                        st.metric("Incidents", int(last_week['Total Incidents']), 
                                 delta=f"{int(incident_change):+d}", delta_color=delta_color)
                    with col_b:
                        delta_color = "inverse" if severity_change <= 0 else "normal"
                        st.metric("Avg Severity", f"{last_week['Avg Severity']:.1f}",
                                 delta=f"{severity_change:+.1f}", delta_color=delta_color)
                    
                    # Trend interpretation
                    if incident_change < 0 and severity_change < 0:
                        st.success("üìâ **Improving:** Both incidents and severity decreasing")
                    elif incident_change > 0 and severity_change > 0:
                        st.error("üìà **Concern:** Both incidents and severity increasing")
                    elif incident_change < 0:
                        st.info("üìä **Mixed:** Fewer incidents but severity unchanged/increasing")
                    else:
                        st.warning("üìä **Monitor:** Patterns require continued observation")
    
    st.markdown("---")
    
    # ================================================================
    # ENHANCED CLINICAL SUMMARY
    # ================================================================
    
    st.markdown("### üß† Clinical Summary & AC-Aligned Recommendations")
    st.caption("Evidence-based interpretation using ABA, Berry Street, CPI, and Australian Curriculum frameworks")
    
    recent = full_df.tail(7)
    risk_score = min(100, int(
        (len(recent) / 7 * 10) +
        (recent["severity"].mean() * 8) +
        (len(full_df[full_df["incident_type"] == "Critical"]) / len(full_df) * 50)
    ))
    risk_level = "LOW" if risk_score < 30 else "MODERATE" if risk_score < 60 else "HIGH"
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.info(f"""
**Key Patterns Identified:**
- Primary behaviour: **{top_beh}**
- Main trigger: **{top_ant}**
- Hotspot location: **{top_loc}**
- Peak time: **{top_session}**
- Risk Level: **{risk_level}** ({risk_score}/100)

**Primary AC Capability Focus:** {ac_info['name'] if ac_info else 'Personal and Social Capability'}
**Berry Street Focus:** Body (regulation) and Relationship (connection) domains are foundation.
        """)
    
    with col2:
        # Risk gauge
        risk_color = '#28A745' if risk_score < 30 else '#FFC107' if risk_score < 60 else '#DC3545'
        st.markdown(f"""
<div style='text-align: center; padding: 1rem; background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%); border-radius: 10px;'>
    <div style='font-size: 3rem; font-weight: bold; color: {risk_color};'>{risk_score}</div>
    <div style='font-size: 0.9rem; color: #666;'>Risk Score</div>
    <div style='font-size: 1.2rem; font-weight: bold; color: {risk_color};'>{risk_level}</div>
</div>
        """, unsafe_allow_html=True)
    
    st.success(f"""
**Evidence-Based Recommendations (AC-Aligned):**

**1. Body Domain (PSC: Self-management)**
- Regulated start before {top_session} session
- Breathing exercises and movement breaks
- *AC Goal: "Express emotions appropriately"*

**2. Relationship Domain (PSC: Social awareness & management)**
- Key adult check-in each morning
- Acknowledgment of feelings before demands
- *AC Goal: "Develop positive relationships"*

**3. Stamina Domain (PSC: Self-management)**
- Teach help-seeking using visual supports
- Practice requesting breaks in calm moments
- *AC Goal: "Persist when faced with challenges"*

**4. SMART Goal (AC-Aligned):**
Over 5 weeks, {student['name']} will use a help-seeking strategy (e.g., break card, signal to adult) in 4 out of 5 opportunities when experiencing triggers related to '{top_ant[:30]}...', with visual and verbal support.

*This goal addresses PSC element: "Work independently and show initiative"*
    """)
    
    st.markdown("---")
    
    # EXPORT
    st.markdown("### üìÑ Export Data & Reports")
    
    col1, col2 = st.columns(2)
    
    with col1:
        csv = full_df.to_csv(index=False)
        st.download_button(
            "üì• Download Raw Data (CSV)",
            csv,
            file_name=f"{student['name'].replace(' ', '_')}_Incident_Data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
            use_container_width=True
        )
    
    with col2:
        with st.spinner("Generating Behaviour Analysis Plan..."):
            docx_file = generate_behaviour_analysis_plan_docx(
                student, full_df, top_ant, top_beh, top_loc, top_session, risk_score, risk_level
            )
        if docx_file:
            st.download_button(
                "üìÑ Behaviour Analysis Plan (Word)",
                docx_file,
                file_name=f"BAP_{student['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("‚¨Ö Back to Students", type="primary", key="back_analysis_bottom", use_container_width=True):
            go_to("program_students", selected_program=student["program"])
    with col2:
        if st.button("üè† Program Landing", key="home_analysis_bottom", use_container_width=True):
            go_to("landing")



def render_admin_portal():
    """Admin portal for managing students and placement dates"""
    if st.session_state.current_user.get("role") != "ADM":
        st.error("‚õî Access Denied: Admin privileges required")
        if st.button("‚¨Ö Back to Landing"):
            go_to("landing")
        return
    
    st.markdown("## üîß Admin Portal")
    
    col1, col2 = st.columns([6, 1])
    with col1:
        if st.button("‚¨Ö Back to Landing", key="back_admin"):
            go_to("landing")
    
    st.markdown("---")
    
    # TABS
    tab1, tab2, tab3 = st.tabs(["üë• Manage Students", "üìä System Overview", "üë®‚Äçüíº Staff Management"])
    
    with tab1:
        st.markdown("### Student Management")
        
        # ADD NEW STUDENT
        with st.expander("‚ûï Add New Student", expanded=False):
            with st.form("add_student_form"):
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    new_first_name = st.text_input("First Name *", placeholder="John")
                    new_last_name = st.text_input("Last Name *", placeholder="Smith")
                
                with col2:
                    new_grade = st.selectbox("Grade *", ["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"])
                    new_edid = st.text_input("EDID *", placeholder="ED123456")
                
                with col3:
                    new_dob = st.date_input("Date of Birth *", value=date(2015, 1, 1), format="DD/MM/YYYY")
                    new_program = st.selectbox("Program *", ["JP", "PY", "SY"])
                
                with col4:
                    new_placement_start = st.date_input("Placement Start Date *", value=date.today(), format="DD/MM/YYYY")
                    new_placement_end = st.date_input("Placement End Date (Optional)", value=None, format="DD/MM/YYYY")
                
                submitted = st.form_submit_button("Add Student", type="primary")
                
                if submitted:
                    if new_first_name and new_last_name and new_grade and new_program and new_edid:
                        full_name = f"{new_first_name} {new_last_name}"
                        new_student = {
                            "id": f"stu_{uuid.uuid4().hex[:8]}",
                            "first_name": new_first_name.strip(),
                            "last_name": new_last_name.strip(),
                            "name": full_name.strip(),
                            "grade": new_grade,
                            "dob": new_dob.isoformat(),
                            "edid": new_edid,
                            "program": new_program,
                            "placement_start": new_placement_start.isoformat(),
                            "placement_end": new_placement_end.isoformat() if new_placement_end else None
                        }
                        
                        # SAVE TO DATABASE FIRST
                        if save_student_to_db(new_student):
                            st.session_state.students.append(new_student)
                            st.success(f"‚úÖ Added {full_name} (EDID: {new_edid}) to {PROGRAM_NAMES[new_program]}")
                            st.rerun()
                        else:
                            st.error("‚ùå Failed to save student to database")
                    else:
                        st.error("Please complete all required fields (First Name, Last Name, Grade, EDID, Program)")
        
        st.markdown("---")
        
        # EXISTING STUDENTS
        st.markdown("### Current Students")
        
        for program in ["JP", "PY", "SY"]:
            st.markdown(f"#### {PROGRAM_NAMES[program]}")
            
            program_students = [s for s in st.session_state.students if s["program"] == program]
            
            if not program_students:
                st.caption("No students in this program")
                continue
            
            for student in program_students:
                with st.container(border=True):
                    col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
                    
                    with col1:
                        st.markdown(f"**{student['name']}**")
                        st.caption(f"Grade {student['grade']}")
                        if student.get('edid'):
                            st.caption(f"üÜî EDID: {student['edid']}")
                    
                    with col2:
                        if student.get('placement_start'):
                            start_date = datetime.fromisoformat(student['placement_start']).strftime('%d/%m/%Y')
                            st.caption(f"üìÖ Start: {start_date}")
                            
                            # Calculate days enrolled
                            start = datetime.fromisoformat(student['placement_start']).date()
                            end = datetime.fromisoformat(student['placement_end']).date() if student.get('placement_end') else date.today()
                            days = (end - start).days
                            st.caption(f"üìä {days} days enrolled")
                        else:
                            st.caption("üìÖ Start: Not set")
                            st.caption("üìä Days: N/A")
                    
                    with col3:
                        if student.get('placement_end'):
                            end_date = datetime.fromisoformat(student['placement_end']).strftime('%d/%m/%Y')
                            st.caption(f"üìÖ End: {end_date}")
                            st.caption("üî¥ Inactive")
                        else:
                            st.caption("üìÖ End: Ongoing")
                            st.caption("üü¢ Active")
                    
                    with col4:
                        if st.button("‚úèÔ∏è", key=f"edit_{student['id']}", help="Edit student"):
                            st.session_state.editing_student = student['id']
                            st.rerun()
                
                # EDIT STUDENT
                if st.session_state.get("editing_student") == student['id']:
                    with st.expander("‚úèÔ∏è Edit Student Details", expanded=True):
                        with st.form(f"edit_form_{student['id']}"):
                            edit_col1, edit_col2 = st.columns(2)
                            
                            with edit_col1:
                                edit_first_name = st.text_input("First Name", 
                                                               value=student.get('first_name', student['name'].split()[0] if student['name'] else ''),
                                                               key=f"edit_first_{student['id']}")
                                edit_last_name = st.text_input("Last Name",
                                                              value=student.get('last_name', ' '.join(student['name'].split()[1:]) if len(student['name'].split()) > 1 else ''),
                                                              key=f"edit_last_{student['id']}")
                                # Use existing date or default to today
                                default_start = datetime.fromisoformat(student['placement_start']).date() if student.get('placement_start') else date.today()
                                edit_start = st.date_input("Placement Start", 
                                                          value=default_start,
                                                          key=f"edit_start_{student['id']}",
                                                          format="DD/MM/YYYY")
                            
                            with edit_col2:
                                edit_edid = st.text_input("EDID", value=student.get('edid', ''), key=f"edit_edid_{student['id']}")
                                edit_grade = st.selectbox("Grade", 
                                                         ["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"],
                                                         index=["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"].index(student['grade']) if student['grade'] in ["R", "Y1", "Y2", "Y3", "Y4", "Y5", "Y6", "Y7", "Y8", "Y9", "Y10", "Y11", "Y12"] else 0,
                                                         key=f"edit_grade_{student['id']}")
                                current_end = datetime.fromisoformat(student['placement_end']).date() if student.get('placement_end') else None
                                edit_end = st.date_input("Placement End (None = Ongoing)",
                                                        value=current_end,
                                                        key=f"edit_end_{student['id']}",
                                                        format="DD/MM/YYYY")
                            
                            col_save, col_cancel = st.columns(2)
                            with col_save:
                                if st.form_submit_button("Save Changes", type="primary"):
                                    student['first_name'] = edit_first_name.strip()
                                    student['last_name'] = edit_last_name.strip()
                                    student['name'] = f"{edit_first_name} {edit_last_name}".strip()
                                    student['edid'] = edit_edid
                                    student['grade'] = edit_grade
                                    student['placement_start'] = edit_start.isoformat()
                                    student['placement_end'] = edit_end.isoformat() if edit_end else None
                                    save_student_to_db(student)
                                    st.session_state.editing_student = None
                                    st.success("‚úÖ Updated")
                                    st.rerun()
                            
                            with col_cancel:
                                if st.form_submit_button("Cancel"):
                                    st.session_state.editing_student = None
                                    st.rerun()
    
    with tab2:
        st.markdown("### System Overview")
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            total_students = len(st.session_state.students)
            active_students = len([s for s in st.session_state.students if not s.get('placement_end')])
            st.metric("Total Students", total_students)
            st.caption(f"{active_students} active")
        
        with col2:
            st.metric("Total Incidents", len(st.session_state.incidents))
        
        with col3:
            critical = len([i for i in st.session_state.incidents if i.get("is_critical")])
            st.metric("Critical Incidents", critical)
        
        with col4:
            st.metric("Staff Members", len(st.session_state.staff))
        
        st.markdown("---")
        
        # Program breakdown
        st.markdown("#### Students by Program")
        for prog in ["JP", "PY", "SY"]:
            count = len([s for s in st.session_state.students if s["program"] == prog])
            active = len([s for s in st.session_state.students if s["program"] == prog and not s.get('placement_end')])
            st.write(f"**{PROGRAM_NAMES[prog]}:** {count} total ({active} active)")
    
    with tab3:
        st.markdown("### Staff Management")
        
        # ADD NEW STAFF
        with st.expander("‚ûï Add New Staff Member", expanded=False):
            with st.form("add_staff_form"):
                col1, col2 = st.columns(2)
                
                with col1:
                    staff_first_name = st.text_input("First Name *", placeholder="Jane")
                    staff_last_name = st.text_input("Last Name *", placeholder="Smith")
                    staff_email = st.text_input("Email Address *", placeholder="jane.smith@school.edu.au", 
                                               help="Will be used as username and for critical incident notifications")
                    staff_password = st.text_input("Initial Password *", type="password", value="demo123")
                
                with col2:
                    staff_role = st.selectbox("Role *", 
                                             ["TSS", "Teacher", "Leader", "ADM"],
                                             help="TSS=Teacher/Support Staff, Leader=Program Leader, ADM=Administrator")
                    staff_program = st.selectbox("Program *", ["JP", "PY", "SY", "All Programs"])
                    staff_phone = st.text_input("Phone Number", placeholder="0412 345 678")
                
                staff_notes = st.text_area("Notes (Optional)", placeholder="Additional information about this staff member")
                
                submit_staff = st.form_submit_button("Add Staff Member", type="primary")
                
                if submit_staff:
                    if staff_first_name and staff_last_name and staff_email and staff_password and staff_role:
                        staff_full_name = f"{staff_first_name} {staff_last_name}"
                        # Check if email already exists
                        if any(s.get("email", "").lower() == staff_email.lower() for s in st.session_state.staff):
                            st.error(f"‚ùå Email {staff_email} already exists")
                        else:
                            new_staff = {
                                "id": f"staff_{uuid.uuid4().hex[:8]}",
                                "first_name": staff_first_name.strip(),
                                "last_name": staff_last_name.strip(),
                                "name": staff_full_name.strip(),
                                "email": staff_email.lower().strip(),
                                "password": staff_password,
                                "role": staff_role,
                                "program": staff_program if staff_program != "All Programs" else None,
                                "phone": staff_phone if staff_phone else None,
                                "notes": staff_notes if staff_notes else None,
                                "receive_critical_emails": True,  # Default to receiving emails
                                "created_date": date.today().isoformat()
                            }
                            # Save to database first
                            if save_staff_to_db(new_staff):
                                st.session_state.staff.append(new_staff)
                                st.success(f"‚úÖ Added {staff_full_name} ({staff_email}) to database")
                                st.rerun()
                            else:
                                st.error("‚ùå Failed to save staff member to database")
                    else:
                        st.error("Please complete all required fields (First Name, Last Name, Email, Password, Role)")
    
        st.markdown("---")
        
        # EXISTING STAFF
        st.markdown("### Current Staff")
        
        # Group by role
        for role in ["ADM", "Leader", "Teacher", "TSS"]:
            role_names = {"ADM": "Administrators", "Leader": "Program Leaders", "Teacher": "Teachers", "TSS": "Support Staff"}
            role_staff = [s for s in st.session_state.staff if s.get("role") == role]
            
            if role_staff:
                st.markdown(f"#### {role_names[role]}")
                
                for staff in role_staff:
                    with st.container(border=True):
                        col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
                        
                        with col1:
                            st.markdown(f"**{staff['name']}**")
                            st.caption(f"üìß {staff['email']}")
                            if staff.get('phone'):
                                st.caption(f"üì± {staff['phone']}")
                        
                        with col2:
                            st.caption(f"**Role:** {staff['role']}")
                            if staff.get('program'):
                                st.caption(f"**Program:** {PROGRAM_NAMES[staff['program']]}")
                            else:
                                st.caption("**Program:** All Programs")
                        
                        with col3:
                            receives_emails = staff.get('receive_critical_emails', True)
                            if receives_emails:
                                st.caption("üì¨ Receives critical alerts")
                            else:
                                st.caption("üì™ No critical alerts")
                            
                            if staff.get('created_date'):
                                st.caption(f"Added: {staff['created_date']}")
                        
                        with col4:
                            if st.button("‚úèÔ∏è", key=f"edit_staff_{staff['id']}", help="Edit staff"):
                                st.session_state.editing_staff = staff['id']
                                st.rerun()
                        
                        # EDIT STAFF
                        if st.session_state.get("editing_staff") == staff['id']:
                            with st.expander("‚úèÔ∏è Edit Staff Details", expanded=True):
                                with st.form(f"edit_staff_form_{staff['id']}"):
                                    edit_col1, edit_col2 = st.columns(2)
                                    
                                    with edit_col1:
                                        edit_first_name = st.text_input("First Name", value=staff.get('first_name', staff['name'].split()[0] if staff['name'] else ''), key=f"edit_staff_first_{staff['id']}")
                                        edit_last_name = st.text_input("Last Name", value=staff.get('last_name', ' '.join(staff['name'].split()[1:]) if len(staff['name'].split()) > 1 else ''), key=f"edit_staff_last_{staff['id']}")
                                        edit_email = st.text_input("Email", value=staff['email'], key=f"edit_staff_email_{staff['id']}")
                                        edit_phone = st.text_input("Phone", value=staff.get('phone', ''), key=f"edit_staff_phone_{staff['id']}")
                                    
                                    with edit_col2:
                                        edit_role = st.selectbox("Role", ["TSS", "Teacher", "Leader", "ADM"],
                                                                index=["TSS", "Teacher", "Leader", "ADM"].index(staff['role']),
                                                                key=f"edit_staff_role_{staff['id']}")
                                     current_program = staff.get('program') if staff.get('program') else "All Programs"
                                        program_list = ["JP", "PY", "SY", "All Programs"]
                                    try:
                                        program_index = program_list.index(current_program)
                                    except ValueError:
                                       program_index = 3  # Default to "All Programs"
                                     current_program = staff.get('program') if staff.get('program') else "All Programs"
                                    program_list = ["JP", "PY", "SY", "All Programs"]
                                    try:
                                        program_index = program_list.index(current_program)
                                    except ValueError:
                                        program_index = 3
                                    
                                    edit_program = st.selectbox("Program", program_list,
                                                               index=program_index,
                                                               key=f"edit_staff_program_{staff['id']}")                                    
                                    edit_notes = st.text_area("Notes", value=staff.get('notes', ''), key=f"edit_staff_notes_{staff['id']}")
                                    
                                    col_save, col_cancel, col_delete = st.columns([1, 1, 1])
                                    with col_save:
                                        if st.form_submit_button("üíæ Save Changes", type="primary"):
                                            staff['first_name'] = edit_first_name.strip()
                                            staff['last_name'] = edit_last_name.strip()
                                            staff['name'] = f"{edit_first_name} {edit_last_name}".strip()
                                            staff['email'] = edit_email.lower().strip()
                                            staff['phone'] = edit_phone if edit_phone else None
                                            staff['role'] = edit_role
                                            staff['program'] = edit_program if edit_program != "All Programs" else None
                                            staff['receive_critical_emails'] = edit_receive_emails
                                            staff['notes'] = edit_notes if edit_notes else None
                                            save_staff_to_db(staff)
                                            st.session_state.editing_staff = None
                                            st.success("‚úÖ Updated")
                                            st.rerun()
                                    
                                    with col_cancel:
                                        if st.form_submit_button("‚ùå Cancel"):
                                            st.session_state.editing_staff = None
                                            st.rerun()
                                    
                                    with col_delete:
                                        if st.form_submit_button("üóëÔ∏è Delete", help="Remove this staff member"):
                                            if staff['role'] != 'ADM' or len([s for s in st.session_state.staff if s.get('role') == 'ADM']) > 1:
                                                st.session_state.staff.remove(staff)
                                                st.session_state.editing_staff = None
                                                st.success("‚úÖ Staff member removed")
                                                st.rerun()
                                            else:
                                                st.error("‚ùå Cannot delete the last administrator")
        
        st.markdown("---")
        
        # EMAIL NOTIFICATION SETTINGS
        st.markdown("### üìß Email Notification Recipients")
        st.caption("Staff members who will receive critical incident notifications:")
        
        email_recipients = [s for s in st.session_state.staff if s.get('receive_critical_emails', True)]
        
        if email_recipients:
            for recipient in email_recipients:
                st.write(f"‚Ä¢ **{recipient['name']}** ({recipient['email']}) - {recipient['role']}")
        else:
            st.warning("‚ö†Ô∏è No staff members set to receive critical incident emails!")






def main():
    init_state()
    
    if not st.session_state.logged_in:
        render_login_page()
        return
    
    page = st.session_state.current_page
    
    if page == "landing": render_landing_page()
    elif page == "program_students": render_program_students_page()
    elif page == "incident_log": render_incident_log_page()
    elif page == "critical_incident": render_critical_incident_page()
    elif page == "student_analysis": render_student_analysis_page()
    elif page == "admin_portal": render_admin_portal()
    else: render_landing_page()

if __name__ == "__main__":
    main()
